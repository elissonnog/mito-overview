#!/usr/bin/env python3
"""Export the markdown preprint draft to a simple .docx with embedded figures."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


IMAGE_RE = re.compile(r"!\[(?P<alt>.*?)\]\((?P<path>.*?)\)")
LINK_RE = re.compile(r"\[(?P<label>[^\]]+)\]\((?P<url>[^)]+)\)")
CODE_RE = re.compile(r"`([^`]+)`")
BOLD_RE = re.compile(r"\*\*(.*?)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)")


def normalize_inline(text: str) -> str:
    text = LINK_RE.sub(lambda m: f"{m.group('label')} ({m.group('url')})", text)
    text = CODE_RE.sub(lambda m: m.group(1), text)
    text = BOLD_RE.sub(lambda m: m.group(1), text)
    text = ITALIC_RE.sub(lambda m: m.group(1), text)
    return text.strip()


def flush_paragraph(document: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    paragraph = document.add_paragraph(normalize_inline(" ".join(buffer)))
    paragraph.style = document.styles["Normal"]
    buffer.clear()


def add_image(document: Document, source_dir: Path, image_path: str) -> None:
    path = (source_dir / image_path).resolve()
    if not path.exists():
        paragraph = document.add_paragraph(f"[Missing figure: {image_path}]")
        paragraph.italic = True
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.2))


def build_doc(source_path: Path, output_path: Path) -> None:
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"

    lines = source_path.read_text(encoding="utf-8").splitlines()
    paragraph_buffer: list[str] = []

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph(document, paragraph_buffer)
            continue

        image_match = IMAGE_RE.fullmatch(stripped)
        if image_match:
            flush_paragraph(document, paragraph_buffer)
            add_image(document, source_path.parent, image_match.group("path"))
            continue

        if stripped.startswith("# "):
            flush_paragraph(document, paragraph_buffer)
            title = normalize_inline(stripped[2:])
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(title)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(16)
            continue

        if stripped.startswith("## "):
            flush_paragraph(document, paragraph_buffer)
            document.add_heading(normalize_inline(stripped[3:]), level=1)
            continue

        if stripped.startswith("### "):
            flush_paragraph(document, paragraph_buffer)
            document.add_heading(normalize_inline(stripped[4:]), level=2)
            continue

        if stripped.startswith("- "):
            flush_paragraph(document, paragraph_buffer)
            paragraph = document.add_paragraph(normalize_inline(stripped[2:]), style="List Bullet")
            paragraph.style.font.name = "Times New Roman"
            continue

        if re.match(r"^\d+\.\s", stripped):
            flush_paragraph(document, paragraph_buffer)
            paragraph = document.add_paragraph(normalize_inline(stripped), style="List Number")
            paragraph.style.font.name = "Times New Roman"
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph(document, paragraph_buffer)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main(argv: list[str]) -> int:
    source = Path(argv[1]) if len(argv) > 1 else Path("paper/preprint_draft.md")
    output = Path(argv[2]) if len(argv) > 2 else Path("paper/mito-overview_manuscript_2026-04-21.docx")
    build_doc(source.resolve(), output.resolve())
    print(output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
