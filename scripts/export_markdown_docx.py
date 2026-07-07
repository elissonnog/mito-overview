#!/usr/bin/env python3
"""Export a markdown document to a simple .docx with images, code blocks, and tables."""

from __future__ import annotations

import re
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


IMAGE_RE = re.compile(r"!\[(?P<alt>.*?)\]\((?P<path>.*?)\)")
LINK_RE = re.compile(r"\[(?P<label>[^\]]+)\]\((?P<url>[^)]+)\)")
CODE_RE = re.compile(r"`([^`]+)`")
BOLD_RE = re.compile(r"\*\*(.*?)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)")
INLINE_MATH_RE = re.compile(r"\\\((.*?)\\\)")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?(\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?\s*$")


def normalize_inline_math(text: str) -> str:
    """Convert small inline LaTeX snippets to readable plain text for Word output."""

    replacements = {
        r"\ge": ">=",
        r"\le": "<=",
        r"\min": "min",
        r"\max": "max",
        r"\cap": "intersect",
        r"\cup": "union",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    text = re.sub(r"\\bar\{([^}]+)\}", r"mean_\1", text)
    text = re.sub(r"\{\\?([A-Za-z]+)\}", r"\1", text)
    text = re.sub(r"_\{([^}]+)\}", r"_\1", text)
    text = re.sub(r"\^\{([^}]+)\}", r"_\1", text)
    text = text.replace("\\", "")
    return text


def normalize_inline(text: str) -> str:
    text = LINK_RE.sub(lambda m: f"{m.group('label')} ({m.group('url')})", text)
    text = INLINE_MATH_RE.sub(lambda m: normalize_inline_math(m.group(1)), text)
    text = CODE_RE.sub(lambda m: m.group(1), text)
    text = BOLD_RE.sub(lambda m: m.group(1), text)
    text = ITALIC_RE.sub(lambda m: m.group(1), text)
    return text.strip()


def flush_paragraph(document: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    paragraph = document.add_paragraph(normalize_inline(" ".join(buffer)))
    paragraph.style = document.styles["Normal"]
    buffer.clear()


def add_image(document: Document, source_dir: Path, image_path: str) -> None:
    path = (source_dir / image_path).resolve()
    if not path.exists():
        paragraph = document.add_paragraph(f"[Missing figure: {image_path}]")
        paragraph.italic = True
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.7))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    table = document.add_table(rows=len(padded), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(padded):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = normalize_inline(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8.5)
                    if r_idx == 0:
                        run.bold = True
            if r_idx == 0:
                set_cell_shading(cell, "D9E2F3")
    document.add_paragraph("")


def parse_table_block(lines: list[str]) -> list[list[str]]:
    parsed: list[list[str]] = []
    for idx, line in enumerate(lines):
        if idx == 1 and TABLE_SEPARATOR_RE.match(line.strip()):
            continue
        stripped = line.strip().strip("|")
        cells = [cell.strip() for cell in stripped.split("|")]
        parsed.append(cells)
    return parsed


def add_code_block(document: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.style = document.styles["Normal"]
        run = paragraph.add_run(line.rstrip("\n"))
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    document.add_paragraph("")


def render_equation_png(text: str, equation_dir: Path, index: int) -> Path | None:
    """Render display math to a PNG when matplotlib is available."""
    try:
        import os

        os.environ.setdefault("MPLCONFIGDIR", str(Path.cwd() / ".mplconfig"))
        Path(os.environ["MPLCONFIGDIR"]).mkdir(parents=True, exist_ok=True)
        import matplotlib.pyplot as plt
    except Exception:
        return None

    out = equation_dir / f"equation_{index:03d}.png"
    fig = plt.figure(figsize=(0.01, 0.01), dpi=220)
    fig.text(0, 0, f"${text}$", fontsize=14, color="black")
    try:
        fig.savefig(out, dpi=220, bbox_inches="tight", pad_inches=0.08, transparent=True)
    except Exception:
        plt.close(fig)
        return None
    plt.close(fig)
    return out


def add_equation_block(document: Document, lines: list[str], equation_dir: Path, equation_index: int) -> None:
    text = " ".join(line.strip() for line in lines if line.strip())
    if not text:
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation_png = render_equation_png(text, equation_dir, equation_index)
    run = paragraph.add_run()
    if equation_png and equation_png.exists():
        if "\\max" in text:
            width_inches = 2.7
        elif "\\frac" in text or "\\sum" in text:
            width_inches = min(3.6, max(2.4, len(text) * 0.028))
        else:
            width_inches = min(3.4, max(1.2, len(text) * 0.026))
        run.add_picture(str(equation_png), width=Inches(width_inches))
    else:
        run.add_text(text)
        run.font.name = "Cambria Math"
        run.font.size = Pt(11)
    document.add_paragraph("")


def build_doc(source_path: Path, output_path: Path) -> None:
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"

    equation_tmp = tempfile.TemporaryDirectory(prefix="mito_overview_docx_equations_")
    equation_dir = Path(equation_tmp.name)
    equation_index = 0
    lines = source_path.read_text(encoding="utf-8").splitlines()
    paragraph_buffer: list[str] = []
    table_buffer: list[str] = []
    code_buffer: list[str] = []
    equation_buffer: list[str] = []
    in_code_block = False
    in_equation_block = False

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        add_table(document, parse_table_block(table_buffer))
        table_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if in_equation_block:
            if stripped == "$$":
                equation_index += 1
                add_equation_block(document, equation_buffer, equation_dir, equation_index)
                equation_buffer = []
                in_equation_block = False
            else:
                equation_buffer.append(line)
            continue

        if in_code_block:
            if stripped.startswith("```"):
                add_code_block(document, code_buffer)
                code_buffer = []
                in_code_block = False
            else:
                code_buffer.append(line)
            continue

        if stripped.startswith("```"):
            flush_paragraph(document, paragraph_buffer)
            flush_table()
            in_code_block = True
            continue

        if stripped == "$$":
            flush_paragraph(document, paragraph_buffer)
            flush_table()
            in_equation_block = True
            continue

        is_table_line = "|" in stripped and not stripped.startswith("![") and not stripped.startswith("[")
        if table_buffer and (is_table_line or TABLE_SEPARATOR_RE.match(stripped)):
            table_buffer.append(line)
            continue
        if table_buffer and not is_table_line and not TABLE_SEPARATOR_RE.match(stripped):
            flush_table()

        if is_table_line and "|" in stripped:
            flush_paragraph(document, paragraph_buffer)
            table_buffer.append(line)
            continue

        if not stripped:
            flush_paragraph(document, paragraph_buffer)
            flush_table()
            continue

        image_match = IMAGE_RE.fullmatch(stripped)
        if image_match:
            flush_paragraph(document, paragraph_buffer)
            flush_table()
            add_image(document, source_path.parent, image_match.group("path"))
            continue

        if stripped.startswith("# "):
            flush_paragraph(document, paragraph_buffer)
            flush_table()
            title = normalize_inline(stripped[2:])
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(title)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(16)
            continue

        if stripped.startswith("## "):
            flush_paragraph(document, paragraph_buffer)
            flush_table()
            document.add_heading(normalize_inline(stripped[3:]), level=1)
            continue

        if stripped.startswith("### "):
            flush_paragraph(document, paragraph_buffer)
            flush_table()
            document.add_heading(normalize_inline(stripped[4:]), level=2)
            continue

        if stripped.startswith("#### "):
            flush_paragraph(document, paragraph_buffer)
            flush_table()
            document.add_heading(normalize_inline(stripped[5:]), level=3)
            continue

        if stripped.startswith("- "):
            flush_paragraph(document, paragraph_buffer)
            flush_table()
            paragraph = document.add_paragraph(normalize_inline(stripped[2:]), style="List Bullet")
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
            continue

        numbered_match = re.match(r"^\d+\.\s+(.*)", stripped)
        if numbered_match:
            flush_paragraph(document, paragraph_buffer)
            flush_table()
            paragraph = document.add_paragraph(normalize_inline(numbered_match.group(1)), style="List Number")
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph(document, paragraph_buffer)
    flush_table()
    if code_buffer:
        add_code_block(document, code_buffer)
    if equation_buffer:
        equation_index += 1
        add_equation_block(document, equation_buffer, equation_dir, equation_index)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    equation_tmp.cleanup()


def main(argv: list[str]) -> int:
    if len(argv) < 2:
        print("Usage: export_markdown_docx.py SOURCE.md [OUTPUT.docx]", file=sys.stderr)
        return 1
    source = Path(argv[1]).resolve()
    if len(argv) > 2:
        output = Path(argv[2]).resolve()
    else:
        output = source.with_suffix(".docx")
    build_doc(source, output)
    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
