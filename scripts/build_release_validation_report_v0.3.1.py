#!/usr/bin/env python3
"""Build the human-readable MitoOverview v0.3.1 release-validation report.

The input is a completed schema-2.0 ``github_release_validation_v1`` packet
plus exact read-only GitHub prepublication metadata. The builder is intentionally
fail-closed: identity drift, a non-PASS case or oracle assertion, an invalid
cross-platform comparison, or a mismatched figure hash prevents report output.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import math
import os
import re
import shutil
import subprocess
import tempfile
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


EXPECTED_SCHEMA = "2.0"
EXPECTED_PROFILE = "github_release_validation_v1"
EXPECTED_VERSION = "v0.3.1"
EXPECTED_PACKAGE_VERSION = "0.3.1"
SCIENTIFIC_PROTOCOL_VERSION = "v0.3.0"
OUTPUT_STEM = "MitoOverview_v0.3.1_release_validation_report"
BUILD_PROVENANCE_NAME = "report_build_provenance.json"
SHA256_RE = re.compile(r"^[0-9a-f]{64}$")
COMMIT_RE = re.compile(r"^[0-9a-f]{40}$")
SAFE_ID_RE = re.compile(r"[^A-Za-z0-9_.-]+")
EMPTY_SHA256 = hashlib.sha256(b"").hexdigest()
PUBLIC_ENVIRONMENT_ROOT = "public_environment"
NETWORK_ISOLATION_PACKET_PATH = f"{PUBLIC_ENVIRONMENT_ROOT}/network_isolation.tsv"
RUNTIME_VERSIONS_PACKET_PATH = f"{PUBLIC_ENVIRONMENT_ROOT}/runtime_versions.json"
GM11906_SOURCE_METADATA_PACKET_PATH = (
    "public_provenance/GM11906_NCBI_source_metadata.json"
)
GM11906_SOURCE_METADATA_SHA256 = (
    "01be488b9dc6bfce0726304be95db4259b1a85a53ac8e620cba4c337842d3185"
)

UBUNTU_CASES_PACKET_PATH = (
    "acceptance/ubuntu_public_validation/artifact/results/cases.tsv"
)
ALLOWED_EMPTY_COMPARISONS = {
    "logs/public/gm11906_repeatability.diff": (
        "cases.tsv",
        "gm11906_repeatability",
        "repeatability",
    ),
    "logs/public/gm12878_repeatability.diff": (
        "cases.tsv",
        "gm12878_repeatability",
        "repeatability",
    ),
    "logs/public/gm11906_visual_structure.diff": (
        "cases.tsv",
        "gm11906_visual_integrity",
        "visual_integrity",
    ),
    "logs/public/gm12878_visual_structure.diff": (
        "cases.tsv",
        "gm12878_visual_integrity",
        "visual_integrity",
    ),
    (
        "acceptance/ubuntu_public_validation/artifact/results/logs/"
        "gm11906_repeatability.diff"
    ): (UBUNTU_CASES_PACKET_PATH, "gm11906_repeatability", "repeatability"),
    (
        "acceptance/ubuntu_public_validation/artifact/results/logs/"
        "gm12878_repeatability.diff"
    ): (UBUNTU_CASES_PACKET_PATH, "gm12878_repeatability", "repeatability"),
    (
        "acceptance/ubuntu_public_validation/artifact/results/logs/"
        "gm11906_visual_structure.diff"
    ): (UBUNTU_CASES_PACKET_PATH, "gm11906_visual_integrity", "visual_integrity"),
    (
        "acceptance/ubuntu_public_validation/artifact/results/logs/"
        "gm12878_visual_structure.diff"
    ): (UBUNTU_CASES_PACKET_PATH, "gm12878_visual_integrity", "visual_integrity"),
}

EXPECTED_RUNTIME_PACKAGES = {
    "mito-overview": EXPECTED_PACKAGE_VERSION,
    "biopython": "1.87",
    "pysam": "0.24.0",
    "pandas": "3.0.3",
    "numpy": "2.5.1",
    "matplotlib": "3.11.0",
    "requests": "2.34.2",
    "pytest": "9.1.1",
    "build": "1.5.0",
    "setuptools": "82.0.1",
    "wheel": "0.47.0",
    "python-docx": "1.2.0",
}

REQUIRED_PACKET_FILES = (
    "run.json",
    "release_identity.json",
    "cases.tsv",
    "claim_evidence_matrix.tsv",
    "module_status_matrix.tsv",
    "resource_usage.tsv",
    "figure_provenance.tsv",
    "table_provenance.tsv",
    "public_data_sources.tsv",
    "manuscript_handoff.tsv",
    "limitations.tsv",
    "environment.txt",
    "filter_profile_results.tsv",
    "inputs.sha256",
    "raw_inputs.tsv",
    "CACHE_SEAL.sha256",
    "public_validation_oracle_v0.3.0.tsv",
    "oracle_assertions.tsv",
    "cross_platform_comparison.tsv",
    "acceptance/cross_platform_public_reproduction.json",
    NETWORK_ISOLATION_PACKET_PATH,
    RUNTIME_VERSIONS_PACKET_PATH,
    GM11906_SOURCE_METADATA_PACKET_PATH,
    "decoded_pixel_hashes/GM11906.tsv",
    "decoded_pixel_hashes/GM12878.tsv",
    "artifacts.sha256",
    "verify_bundle.sh",
)

REQUIRED_CASE_IDS = frozenset(
    {
        "unit_known_answer",
        "public_validation_matrix",
        "gm11906_default_run1",
        "gm12878_default_run1",
        "gm11906_repeatability",
        "gm12878_repeatability",
        "gm11906_visual_integrity",
        "gm12878_visual_integrity",
        "cross_platform_public_reproduction",
        "fresh_clone_candidate_commit",
        "github_actions_linux_candidate_commit",
        "github_actions_macos_candidate_commit",
        "github_actions_macos_arm64_candidate_commit",
    }
)

REQUIRED_RESOURCE_CASE_IDS = frozenset(
    {
        "fresh_clone_candidate_commit",
        "package_build",
        "unit_known_answer",
        "cli_step_listing",
        "strict_generic_dry_run",
        "synthetic_longread_smoke",
        "synthetic_shortread_smoke",
        "synthetic_longread_nomethyl_smoke",
        "standalone_minimal_smoke",
        "public_cache_prepare",
        "public_validation_matrix",
    }
)
RESOURCE_CASE_THREAD_SETTINGS = {
    "fresh_clone_candidate_commit": "mixed",
    "package_build": "not_applicable",
    "unit_known_answer": "mixed",
    "cli_step_listing": "not_applicable",
    "strict_generic_dry_run": "4",
    "synthetic_longread_smoke": "1",
    "synthetic_shortread_smoke": "1",
    "synthetic_longread_nomethyl_smoke": "1",
    "standalone_minimal_smoke": "4",
    "public_cache_prepare": "not_applicable",
    "public_validation_matrix": "4",
}

ALLOWED_MODULE_STATES = frozenset(
    {
        "ok",
        "not_configured",
        "not_applicable",
        "not_evaluable",
        "unavailable",
        "failed",
    }
)

EVIDENCE_COLUMNS = {
    "cases.tsv": (
        "case_id",
        "category",
        "input_available",
        "expected_available",
        "verdict",
        "detail",
    ),
    "claim_evidence_matrix.tsv": (
        "claim_id",
        "bounded_claim",
        "evidence",
        "limitation",
    ),
    "module_status_matrix.tsv": (
        "dataset",
        "case_id",
        "module",
        "status",
        "reason_code",
        "source_table",
    ),
    "resource_usage.tsv": (
        "measurement_id",
        "case_id",
        "candidate_commit",
        "command_path",
        "command_sha256",
        "packaged_command_sha256",
        "log_path",
        "log_sha256",
        "packaged_log_sha256",
        "wall_seconds",
        "user_cpu_seconds",
        "system_cpu_seconds",
        "max_rss_kb",
        "broad_declared_input_inventory_file_count",
        "broad_declared_input_inventory_bytes",
        "changed_or_new_output_inventory_file_count",
        "changed_or_new_output_inventory_bytes",
        "broad_declared_input_inventory_scope",
        "changed_or_new_output_inventory_scope",
        "io_measurement_method",
        "threads",
        "platform",
        "measurement_status",
        "reason",
    ),
    "figure_provenance.tsv": (
        "figure_id",
        "dataset",
        "case_id",
        "packet_path",
        "sha256",
        "bytes",
        "width",
        "height",
        "visual_status",
        "source_inventory",
    ),
    "table_provenance.tsv": (
        "table_id",
        "dataset",
        "case_id",
        "packet_path",
        "sha256",
        "rows",
        "columns",
        "purpose",
    ),
    "public_data_sources.tsv": (
        "dataset",
        "run_accession",
        "study_accession",
        "sample_accession",
        "cell_line",
        "platform",
        "instrument_model",
        "library_strategy",
        "fastq_url",
        "fastq_md5",
        "fastq_sha256",
        "fastq_bytes",
        "metadata_recorded_utc",
        "role",
        "redistribution",
    ),
    "manuscript_handoff.tsv": (
        "result_id",
        "dataset",
        "metric",
        "value",
        "unit",
        "source_table",
        "claim_boundary",
    ),
    "limitations.tsv": (
        "limitation_id",
        "scope",
        "limitation",
        "release_effect",
    ),
    "oracle_assertions.tsv": (
        "assertion_id",
        "verdict",
        "expected",
        "observed",
        "detail",
    ),
    "cross_platform_comparison.tsv": (
        "evidence_type",
        "relative_path",
        "macos_sha256",
        "ubuntu_sha256",
        "verdict",
        "comparison",
    ),
    "filter_profile_results.tsv": (
        "case_id",
        "dataset",
        "profile",
        "min_base_quality",
        "min_mapping_quality",
        "min_read_mean_quality",
        "candidate_sites",
        "accepted_observations",
        "excluded_observations",
        "m8344_A_G_present",
        "m8344_A_G_alt_allele_fraction",
    ),
    "public_validation_oracle_v0.3.0.tsv": (
        "dataset",
        "profile",
        "candidate_sites",
        "accepted_observations",
        "excluded_observations",
        "m8344_alt_fraction",
        "summary_tsv_count",
        "html_count",
        "png_count",
    ),
}

RAW_INPUT_COLUMNS = (
    "schema_version",
    "dataset_id",
    "run_accession",
    "sample_accession",
    "sample_alias",
    "sample_title",
    "source_sample_id",
    "library_strategy",
    "library_unit",
    "source_record_url",
    "filename",
    "bytes",
    "md5",
    "sha256",
    "fastq_records",
    "url",
)


class ReportValidationError(ValueError):
    """Raised when release evidence is incomplete or inconsistent."""


@dataclass(frozen=True)
class Heading:
    level: int
    text: str


@dataclass(frozen=True)
class Paragraph:
    text: str


@dataclass(frozen=True)
class Formula:
    text: str


@dataclass(frozen=True)
class CodeBlock:
    text: str


@dataclass(frozen=True)
class TableBlock:
    title: str
    columns: tuple[tuple[str, str], ...]
    rows: tuple[dict[str, str], ...]


@dataclass(frozen=True)
class FigureBlock:
    number: int
    dataset: str
    case_id: str
    title: str
    asset_path: Path
    asset_relative: str
    packet_path: str
    sha256: str
    width: int
    height: int
    source_inventory: str


@dataclass(frozen=True)
class PageBreak:
    pass


Block = Heading | Paragraph | Formula | CodeBlock | TableBlock | FigureBlock | PageBreak


@dataclass
class ReportEvidence:
    packet_root: Path
    run: dict[str, object]
    release: dict[str, object]
    publication: dict[str, object]
    tables: dict[str, list[dict[str, str]]]
    raw_inputs: list[dict[str, str]]
    environment_text: str
    runtime_versions: dict[str, object]
    network_isolation: dict[str, str]
    cross_platform: dict[str, object]
    public_source_metadata: dict[str, object]
    figures: list[dict[str, object]]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--packet-root",
        type=Path,
        required=True,
        help="Extracted schema-2.0 validation packet root.",
    )
    parser.add_argument(
        "--publication-json",
        type=Path,
        help=(
            "Read-only github_prepublication.json captured before the report "
            "becomes a release asset."
        ),
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        help="Destination for the Markdown, DOCX, and report figure assets.",
    )
    parser.add_argument(
        "--preflight-packet",
        action="store_true",
        help=(
            "Validate the complete packet without publication metadata or report "
            "output."
        ),
    )
    parser.add_argument(
        "--overwrite",
        action="store_true",
        help="Replace a report previously generated in the output directory.",
    )
    parser.add_argument(
        "--emit-pdf",
        action="store_true",
        help=(
            "Optionally hand the generated DOCX to LibreOffice for PDF conversion. "
            "This does not replace the separate rendered-page QA workflow."
        ),
    )
    args = parser.parse_args()
    if args.preflight_packet:
        incompatible = (
            args.publication_json is not None
            or args.output_dir is not None
            or args.overwrite
            or args.emit_pdf
        )
        if incompatible:
            parser.error(
                "--preflight-packet cannot be combined with report-output options"
            )
    elif args.publication_json is None or args.output_dir is None:
        parser.error(
            "report generation requires --publication-json and --output-dir"
        )
    return args


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def file_record(path: Path, *, name: str | None = None) -> dict[str, object]:
    """Return a portable content-identity record for one report input/output."""

    require_plain_file(path, name or path.name)
    return {
        "name": name or path.name,
        "bytes": path.stat().st_size,
        "sha256": sha256_file(path),
    }


def require_plain_file(path: Path, label: str) -> None:
    if path.is_symlink() or not path.is_file() or path.stat().st_size == 0:
        raise ReportValidationError(f"Missing, empty, or symlinked {label}: {path}")


def packet_path(root: Path, relative: str) -> Path:
    candidate = root / relative
    if candidate.is_symlink():
        raise ReportValidationError(f"Packet evidence must not be a symlink: {relative}")
    resolved_root = root.resolve()
    try:
        candidate.resolve().relative_to(resolved_root)
    except ValueError as error:
        raise ReportValidationError(
            f"Packet evidence escapes the packet root: {relative}"
        ) from error
    return candidate


def read_json_object(path: Path, label: str) -> dict[str, object]:
    require_plain_file(path, label)
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (UnicodeDecodeError, json.JSONDecodeError) as error:
        raise ReportValidationError(f"Invalid {label}: {path}: {error}") from error
    if not isinstance(value, dict):
        raise ReportValidationError(f"{label} must be a JSON object: {path}")
    return value


def read_tsv(
    path: Path,
    required_columns: Sequence[str] = (),
    *,
    allow_empty: bool = False,
) -> list[dict[str, str]]:
    require_plain_file(path, "TSV evidence")
    with path.open(encoding="utf-8", newline="") as handle:
        reader = csv.DictReader(handle, delimiter="\t")
        fieldnames = reader.fieldnames or []
        if len(fieldnames) != len(set(fieldnames)):
            raise ReportValidationError(f"Duplicate TSV columns in {path}")
        missing = [name for name in required_columns if name not in fieldnames]
        if missing:
            raise ReportValidationError(
                f"Missing TSV columns in {path}: {', '.join(missing)}"
            )
        rows = [dict(row) for row in reader]
    if not rows and not allow_empty:
        raise ReportValidationError(f"TSV evidence contains no rows: {path}")
    return rows


def read_key_value_tsv(path: Path, label: str) -> dict[str, str]:
    require_plain_file(path, label)
    with path.open(encoding="utf-8", newline="") as handle:
        reader = csv.DictReader(handle, delimiter="\t")
        if tuple(reader.fieldnames or ()) != ("field", "value"):
            raise ReportValidationError(
                f"{label} must contain exactly field and value columns: {path}"
            )
        rows = list(reader)
    if not rows:
        raise ReportValidationError(f"{label} contains no evidence rows: {path}")
    values: dict[str, str] = {}
    for row in rows:
        field = row.get("field", "")
        value = row.get("value", "")
        if not field or not value:
            raise ReportValidationError(f"{label} contains an empty field or value")
        if field in values:
            raise ReportValidationError(f"{label} contains duplicate field {field!r}")
        values[field] = value
    return values


def parse_sha256_manifest(path: Path) -> dict[str, str]:
    require_plain_file(path, "SHA-256 manifest")
    records: dict[str, str] = {}
    for line_number, line in enumerate(path.read_text(encoding="utf-8").splitlines(), 1):
        match = re.fullmatch(r"([0-9a-f]{64})  (.+)", line)
        if not match:
            raise ReportValidationError(
                f"Malformed SHA-256 manifest line {path}:{line_number}"
            )
        digest, relative = match.groups()
        relative_path = Path(relative)
        if relative_path.is_absolute() or ".." in relative_path.parts:
            raise ReportValidationError(f"Unsafe manifest path in {path}: {relative}")
        if relative in records:
            raise ReportValidationError(f"Duplicate manifest path in {path}: {relative}")
        records[relative] = digest
    if not records:
        raise ReportValidationError(f"SHA-256 manifest contains no records: {path}")
    return records


def validate_runtime_versions(
    path: Path, expected_package_version: str = EXPECTED_PACKAGE_VERSION
) -> dict[str, object]:
    runtime = read_json_object(path, "public runtime-version evidence")
    required = {
        "schema_version",
        "platform_id",
        "system",
        "machine",
        "python",
        "python_executable",
        "mito_overview_module",
        "packages",
        "samtools",
        "htslib",
        "minimap2",
        "bwa",
        "threads",
        "installed_distribution_required",
    }
    missing = sorted(required - set(runtime))
    if missing:
        raise ReportValidationError(
            "Public runtime-version evidence is missing: " + ", ".join(missing)
        )
    if runtime["schema_version"] != "1.0":
        raise ReportValidationError("Public runtime-version schema must be 1.0")

    platform_contract = {
        "linux-64": ("Linux", "x86_64", "linux_unshare_network_namespace"),
        "osx-64": ("Darwin", "x86_64", "macos_sandbox_exec_deny_network"),
        "osx-arm64": ("Darwin", "arm64", "macos_sandbox_exec_deny_network"),
    }
    platform_id = runtime["platform_id"]
    if platform_id not in platform_contract:
        raise ReportValidationError(
            f"Public report runtime platform is unsupported: {platform_id!r}"
        )
    if (runtime["system"], runtime["machine"]) != platform_contract[platform_id][:2]:
        raise ReportValidationError(
            "Public runtime platform_id does not match system and machine"
        )
    if runtime["python"] != "3.12.13":
        raise ReportValidationError(
            f"Public runtime Python mismatch: {runtime['python']!r} != '3.12.13'"
        )
    expected_runtime_packages = dict(EXPECTED_RUNTIME_PACKAGES)
    expected_runtime_packages["mito-overview"] = expected_package_version
    if runtime["packages"] != expected_runtime_packages:
        raise ReportValidationError(
            "Public runtime package versions do not match the locked environment"
        )
    expected_tools = {
        "samtools": "samtools 1.23.1",
        "htslib": "Using htslib 1.23.1",
        "minimap2": "2.31-r1302",
        "bwa": "0.7.19-r1273",
    }
    for field, expected in expected_tools.items():
        if runtime[field] != expected:
            raise ReportValidationError(
                f"Public runtime {field} mismatch: {runtime[field]!r} != {expected!r}"
            )
    if runtime["threads"] != 4:
        raise ReportValidationError("Public runtime must record exactly four threads")
    if runtime["installed_distribution_required"] is not True:
        raise ReportValidationError(
            "Public runtime must require the installed mito-overview distribution"
        )
    for field in ("python_executable", "mito_overview_module"):
        if not isinstance(runtime[field], str) or not runtime[field]:
            raise ReportValidationError(f"Public runtime {field} is empty or invalid")
    module_path = str(runtime["mito_overview_module"]).replace("\\", "/")
    if "site-packages/mito_overview/__init__.py" not in module_path:
        raise ReportValidationError(
            "Public runtime did not resolve mito-overview from an installed distribution"
        )
    return runtime


def validate_network_isolation(
    path: Path, runtime: dict[str, object]
) -> dict[str, str]:
    evidence = read_key_value_tsv(path, "public OS-level network-isolation evidence")
    required = {
        "schema_version",
        "platform",
        "isolation_method",
        "isolation_scope",
        "parent_loopback_control",
        "isolated_loopback_probe",
        "probe_target",
        "probe_error",
        "invoking_uid",
        "invoking_gid",
        "child_uid",
        "child_gid",
        "network_isolation_verdict",
    }
    if set(evidence) != required:
        missing = sorted(required - set(evidence))
        extra = sorted(set(evidence) - required)
        raise ReportValidationError(
            "Public network-isolation evidence fields differ from the v0.3.0 contract: "
            f"missing={missing}, extra={extra}"
        )
    expected_platform = f"{runtime['system']}/{runtime['machine']}"
    platform_contract = {
        "linux-64": "linux_unshare_network_namespace",
        "osx-64": "macos_sandbox_exec_deny_network",
        "osx-arm64": "macos_sandbox_exec_deny_network",
    }
    expected = {
        "schema_version": "1.0",
        "platform": expected_platform,
        "isolation_method": platform_contract[str(runtime["platform_id"])],
        "isolation_scope": "process_tree",
        "parent_loopback_control": "reachable",
        "isolated_loopback_probe": "blocked",
        "probe_target": "parent_loopback_listener",
        "network_isolation_verdict": "PASS",
    }
    for field, value in expected.items():
        if evidence[field] != value:
            raise ReportValidationError(
                f"Public network-isolation {field} mismatch: "
                f"{evidence[field]!r} != {value!r}"
            )
    if evidence["probe_error"] == "connection_succeeded":
        raise ReportValidationError(
            "Public network-isolation probe reports a successful connection"
        )
    for role in ("uid", "gid"):
        invoking = evidence[f"invoking_{role}"]
        child = evidence[f"child_{role}"]
        if not invoking.isdigit() or child != invoking:
            raise ReportValidationError(
                f"Public network-isolation child {role} does not match the invoking identity"
            )
    return evidence


def load_manifested_case_rows(
    root: Path, manifest: dict[str, str], relative: str
) -> list[dict[str, str]]:
    if relative not in manifest:
        raise ReportValidationError(
            f"Comparison case table is absent from artifacts.sha256: {relative}"
        )
    path = packet_path(root, relative)
    require_plain_file(path, f"comparison case table {relative}")
    observed = sha256_file(path)
    if observed != manifest[relative]:
        raise ReportValidationError(
            f"Packet artifact hash mismatch for {relative}: "
            f"{observed} != {manifest[relative]}"
        )
    return read_tsv(path, EVIDENCE_COLUMNS["cases.tsv"])


def validate_empty_comparison(
    root: Path,
    manifest: dict[str, str],
    relative: str,
    case_rows: dict[str, list[dict[str, str]]],
) -> None:
    case_table, case_id, category = ALLOWED_EMPTY_COMPARISONS[relative]
    artifact = packet_path(root, relative)
    if artifact.is_symlink() or not artifact.is_file():
        raise ReportValidationError(
            f"Empty comparison evidence must be a regular non-symlink file: {relative}"
        )
    if artifact.stat().st_size != 0:
        raise ReportValidationError(
            f"PASS comparison evidence must be exactly zero bytes: {relative}"
        )
    if manifest.get(relative) != EMPTY_SHA256:
        raise ReportValidationError(
            f"Empty comparison manifest digest is not canonical: {relative}"
        )
    if sha256_file(artifact) != EMPTY_SHA256:
        raise ReportValidationError(
            f"Empty comparison content digest is not canonical: {relative}"
        )
    matches = [row for row in case_rows[case_table] if row["case_id"] == case_id]
    if len(matches) != 1:
        raise ReportValidationError(
            f"Empty comparison requires exactly one case row: {relative} -> {case_id}"
        )
    row = matches[0]
    if row["category"] != category:
        raise ReportValidationError(
            f"Empty comparison case category mismatch: {relative} -> {case_id}"
        )
    if row["verdict"] != "PASS":
        raise ReportValidationError(
            f"Empty comparison case is not PASS: {relative} -> {case_id}"
        )
    if row["input_available"] != "1" or row["expected_available"] != "1":
        raise ReportValidationError(
            f"Empty comparison PASS case lacks available evidence: {relative} -> {case_id}"
        )


def validate_artifact_manifest(root: Path) -> None:
    manifest = parse_sha256_manifest(root / "artifacts.sha256")
    missing_comparisons = sorted(set(ALLOWED_EMPTY_COMPARISONS) - set(manifest))
    if missing_comparisons:
        raise ReportValidationError(
            "Required zero-difference evidence is absent from artifacts.sha256: "
            + ", ".join(missing_comparisons)
        )
    case_rows = {
        relative: load_manifested_case_rows(root, manifest, relative)
        for relative in {value[0] for value in ALLOWED_EMPTY_COMPARISONS.values()}
    }
    for relative, expected in manifest.items():
        artifact = packet_path(root, relative)
        if relative in ALLOWED_EMPTY_COMPARISONS:
            validate_empty_comparison(root, manifest, relative, case_rows)
            continue
        require_plain_file(artifact, f"manifested artifact {relative}")
        observed = sha256_file(artifact)
        if observed != expected:
            raise ReportValidationError(
                f"Packet artifact hash mismatch for {relative}: {observed} != {expected}"
            )
    required_manifested = set(REQUIRED_PACKET_FILES) - {"artifacts.sha256"}
    missing = sorted(required_manifested - set(manifest))
    if missing:
        raise ReportValidationError(
            "Required report evidence is absent from artifacts.sha256: "
            + ", ".join(missing)
        )
    packet_entries = list(root.rglob("*"))
    symlinks = sorted(
        path.relative_to(root).as_posix()
        for path in packet_entries
        if path.is_symlink()
    )
    if symlinks:
        raise ReportValidationError(
            "Packet contains symlinked entries: " + ", ".join(symlinks)
        )
    actual = {
        path.relative_to(root).as_posix()
        for path in packet_entries
        if path.is_file() and path.relative_to(root).as_posix() != "artifacts.sha256"
    }
    unmanifested = sorted(actual - set(manifest))
    if unmanifested:
        raise ReportValidationError(
            "Packet contains unmanifested files: " + ", ".join(unmanifested)
        )


def identity_value(objects: Sequence[tuple[str, dict[str, object]]], field: str) -> object:
    values = [(label, value.get(field)) for label, value in objects]
    first = values[0][1]
    if any(value != first for _, value in values[1:]):
        detail = ", ".join(f"{label}={value!r}" for label, value in values)
        raise ReportValidationError(f"Release identity mismatch for {field}: {detail}")
    return first


def validate_publication(
    publication: dict[str, object], run: dict[str, object], release: dict[str, object]
) -> None:
    required = {
        "schema_version",
        "release_version",
        "git_commit",
        "repository",
        "release_tag",
        "github_release_url",
        "github_actions_run_id",
        "publication_state",
        "verification_state",
        "verified",
        "tag_ref",
        "tag_object",
        "hosting_protection",
        "release",
    }
    missing = sorted(required - set(publication))
    if missing:
        raise ReportValidationError(
            "GitHub publication metadata is missing: " + ", ".join(missing)
        )
    if publication["schema_version"] != "1.0":
        raise ReportValidationError("GitHub publication metadata schema must be 1.0")
    if publication["release_tag"] != EXPECTED_VERSION:
        raise ReportValidationError(f"GitHub release tag must be {EXPECTED_VERSION}")
    if publication["publication_state"] != "prepublication":
        raise ReportValidationError(
            "release-asset reports require a prepublication identity receipt"
        )
    if publication["verified"] is not True:
        raise ReportValidationError("GitHub prepublication receipt is not verified")
    if (
        publication["verification_state"] != "verified_prepublication_identity"
        or publication.get("github_api_read_only") is not True
        or publication.get("mutations_performed") is not False
        or publication.get("asset_publication_verified") is not False
        or publication.get("release_absent") is not True
    ):
        raise ReportValidationError(
            "GitHub prepublication receipt is not read-only or has an invalid state"
        )

    commit = str(run["git_commit"])
    tag_ref = publication["tag_ref"]
    tag_object = publication["tag_object"]
    if not isinstance(tag_ref, dict) or not isinstance(tag_object, dict):
        raise ReportValidationError("GitHub publication receipt lacks annotated-tag identity")
    tag_object_sha = str(tag_ref.get("object_sha", ""))
    if (
        tag_ref.get("ref") != f"refs/tags/{EXPECTED_VERSION}"
        or tag_ref.get("object_type") != "tag"
        or COMMIT_RE.fullmatch(tag_object_sha) is None
        or tag_object.get("tag") != EXPECTED_VERSION
        or tag_object.get("tag_object_sha") != tag_object_sha
        or tag_object.get("target_type") != "commit"
        or tag_object.get("peeled_target_sha") != commit
    ):
        raise ReportValidationError("GitHub publication annotated-tag identity is invalid")

    hosting = publication["hosting_protection"]
    hosting_query_valid = (
        isinstance(hosting, dict)
        and hosting.get("supported") is True
        and hosting.get("enabled") in {True, False}
        and hosting.get("fallback_active") in {None, False}
        and hosting.get("fallback") in {None}
        and hosting.get("reason") in {"queried", "not_enabled"}
    )
    if not hosting_query_valid:
        raise ReportValidationError("GitHub prepublication hosting-state query is invalid")

    release_record = publication["release"]
    if not isinstance(release_record, dict):
        raise ReportValidationError("GitHub prepublication release record is malformed")
    if (
        release_record.get("id") is not None
        or release_record.get("tag_name") != EXPECTED_VERSION
        or release_record.get("target_commitish") != commit
        or release_record.get("url")
        != f"{str(run['repository']).rstrip('/')}/releases/tag/{EXPECTED_VERSION}"
        or release_record.get("draft") is not None
        or release_record.get("immutable") is not None
        or release_record.get("published_at") is not None
    ):
        raise ReportValidationError("GitHub prepublication release identity is invalid")

    repository = str(run["repository"]).rstrip("/")
    expected_url = f"{repository}/releases/tag/{EXPECTED_VERSION}"
    if publication["github_release_url"] != expected_url:
        raise ReportValidationError(
            "GitHub release URL does not match repository and tag: "
            f"{publication['github_release_url']!r} != {expected_url!r}"
        )
    identity_value(
        (("run", run), ("release", release), ("publication", publication)),
        "release_version",
    )
    identity_value(
        (("run", run), ("release", release), ("publication", publication)),
        "git_commit",
    )
    identity_value(
        (("run", run), ("release", release), ("publication", publication)),
        "repository",
    )
    identity_value(
        (("run", run), ("publication", publication)), "github_actions_run_id"
    )


def validate_publication_assets(
    publication: dict[str, object], redownload_field: str
) -> None:
    if publication.get("asset_upload_verified") is not True:
        raise ReportValidationError("GitHub release asset upload is not verified")
    manifest = publication.get("local_asset_manifest")
    remote_assets = publication.get("remote_assets")
    redownload = publication.get(redownload_field)
    if (
        not isinstance(manifest, dict)
        or manifest.get("manifest_name") != "SHA256SUMS"
        or SHA256_RE.fullmatch(str(manifest.get("sha256sums_sha256", ""))) is None
        or not isinstance(manifest.get("assets"), list)
        or not isinstance(remote_assets, list)
        or not isinstance(redownload, dict)
        or redownload.get("verified") is not True
        or redownload.get("method") != "authenticated_redownload_sha256"
        or not isinstance(redownload.get("assets"), list)
    ):
        raise ReportValidationError("GitHub release asset-verification evidence is malformed")

    def asset_map(values: object, *, remote: bool) -> dict[str, tuple[int, str]]:
        if not isinstance(values, list) or not values:
            raise ReportValidationError("GitHub release asset inventory is empty")
        result: dict[str, tuple[int, str]] = {}
        for item in values:
            if not isinstance(item, dict):
                raise ReportValidationError("GitHub release asset inventory is malformed")
            name = item.get("name")
            size = item.get("size")
            digest = item.get("verified_sha256" if remote else "sha256")
            if (
                not isinstance(name, str)
                or not name
                or name in result
                or not isinstance(size, int)
                or size < 0
                or SHA256_RE.fullmatch(str(digest)) is None
            ):
                raise ReportValidationError("GitHub release asset inventory is invalid")
            result[name] = (size, str(digest))
        return result

    local = asset_map(manifest["assets"], remote=False)
    remote = asset_map(remote_assets, remote=True)
    downloaded = asset_map(redownload["assets"], remote=False)
    if local != remote or local != downloaded:
        raise ReportValidationError("GitHub release asset verification inventories differ")


def validate_identity(
    run: dict[str, object],
    release: dict[str, object],
    publication: dict[str, object] | None,
    *,
    expected_release_version: str = EXPECTED_VERSION,
    expected_package_version: str = EXPECTED_PACKAGE_VERSION,
) -> None:
    required_run = {
        "schema_version",
        "validation_profile",
        "release_version",
        "git_commit",
        "repository",
        "github_actions_run_id",
        "generated_utc",
        "case_count",
        "verdict_counts",
        "claim_scope",
        "diagnostic_validation_claimed",
    }
    missing_run = sorted(required_run - set(run))
    if missing_run:
        raise ReportValidationError(
            "run.json is missing release identity fields: " + ", ".join(missing_run)
        )
    required_release = {
        "schema_version",
        "validation_profile",
        "release_version",
        "package_name",
        "package_version",
        "git_commit",
        "repository",
    }
    missing_release = sorted(required_release - set(release))
    if missing_release:
        raise ReportValidationError(
            "release_identity.json is missing fields: " + ", ".join(missing_release)
        )
    for label, value in (("run.json", run), ("release_identity.json", release)):
        if value.get("schema_version") != EXPECTED_SCHEMA:
            raise ReportValidationError(f"{label} schema must be {EXPECTED_SCHEMA}")
        if value.get("validation_profile") != EXPECTED_PROFILE:
            raise ReportValidationError(
                f"{label} validation profile must be {EXPECTED_PROFILE}"
            )
        if value.get("release_version") != expected_release_version:
            raise ReportValidationError(
                f"{label} release must be {expected_release_version}"
            )

    commit = run.get("git_commit")
    if not isinstance(commit, str) or not COMMIT_RE.fullmatch(commit):
        raise ReportValidationError("Release commit must be an exact 40-character SHA")
    if (
        release.get("package_name") != "mito-overview"
        or release.get("package_version") != expected_package_version
    ):
        raise ReportValidationError(
            "Release package identity must be "
            f"mito-overview {expected_package_version}"
        )
    repository = run.get("repository")
    if not isinstance(repository, str) or not re.fullmatch(
        r"https://github\.com/[A-Za-z0-9_.-]+/[A-Za-z0-9_.-]+", repository
    ):
        raise ReportValidationError("Repository must be a canonical public GitHub HTTPS URL")
    run_id = run.get("github_actions_run_id")
    if isinstance(run_id, bool) or not isinstance(run_id, int) or run_id <= 0:
        raise ReportValidationError("github_actions_run_id must be a positive integer")
    try:
        datetime.fromisoformat(str(run["generated_utc"]).replace("Z", "+00:00"))
    except ValueError as error:
        raise ReportValidationError("run.json generated_utc is not ISO-8601") from error
    if run.get("diagnostic_validation_claimed") is not False:
        raise ReportValidationError(
            "The release report cannot assert diagnostic validation"
        )
    expected_scope = "reproducible mode-gated mtDNA reporting workflow/resource"
    if run.get("claim_scope") != expected_scope:
        raise ReportValidationError(
            f"Unexpected release claim scope: {run.get('claim_scope')!r}"
        )
    if publication is not None:
        validate_publication(publication, run, release)


def validate_cases(run: dict[str, object], rows: list[dict[str, str]]) -> None:
    identifiers = [row["case_id"] for row in rows]
    if len(identifiers) != len(set(identifiers)):
        raise ReportValidationError("cases.tsv contains duplicate case identifiers")
    missing = sorted(REQUIRED_CASE_IDS - set(identifiers))
    if missing:
        raise ReportValidationError(
            "Required validation cases are missing: " + ", ".join(missing)
        )
    for row in rows:
        if row["verdict"] != "PASS":
            raise ReportValidationError(
                f"Non-PASS validation case {row['case_id']}: {row['verdict']}"
            )
        if row["input_available"] != "1" or row["expected_available"] != "1":
            raise ReportValidationError(
                f"PASS case lacks input or expected evidence: {row['case_id']}"
            )
    if run.get("case_count") != len(rows):
        raise ReportValidationError(
            f"run.json case_count does not match cases.tsv: {run.get('case_count')} != {len(rows)}"
        )
    counts = run.get("verdict_counts")
    if not isinstance(counts, dict) or counts.get("PASS") != len(rows):
        raise ReportValidationError("run.json verdict_counts does not match cases.tsv")
    nonpass = sum(int(counts.get(name, 0)) for name in ("FAIL", "SKIP", "XFAIL", "BLOCKED"))
    if nonpass:
        raise ReportValidationError("run.json records a non-PASS validation verdict")


def validate_inputs(root: Path, rows: list[dict[str, str]]) -> None:
    input_hashes = parse_sha256_manifest(root / "inputs.sha256")
    seen: set[str] = set()
    for row in rows:
        filename = row["filename"]
        digest = row["sha256"]
        if filename in seen:
            raise ReportValidationError(f"Duplicate frozen input filename: {filename}")
        seen.add(filename)
        if Path(filename).name != filename:
            raise ReportValidationError(f"Frozen input filename is not a basename: {filename}")
        if not SHA256_RE.fullmatch(digest):
            raise ReportValidationError(f"Invalid frozen input SHA-256: {filename}")
        if row["schema_version"] != "1.0":
            raise ReportValidationError(f"Invalid raw input schema: {filename}")
        if not re.fullmatch(r"[0-9a-f]{32}", row["md5"]):
            raise ReportValidationError(f"Invalid frozen input MD5: {filename}")
        if input_hashes.get(filename) != digest:
            raise ReportValidationError(
                f"inputs.sha256 does not bind raw_inputs.tsv for {filename}"
            )
        try:
            if int(row["bytes"]) <= 0 or int(row["fastq_records"]) <= 0:
                raise ValueError
        except ValueError as error:
            raise ReportValidationError(
                f"Frozen input has invalid byte or FASTQ record count: {filename}"
            ) from error
        if not row["url"].startswith("https://"):
            raise ReportValidationError(f"Frozen input URL is not HTTPS: {filename}")
    if set(input_hashes) != seen:
        raise ReportValidationError(
            "inputs.sha256 and raw_inputs.tsv contain different input inventories"
        )
    cache_seal = parse_sha256_manifest(root / "CACHE_SEAL.sha256")
    if cache_seal != {"raw_inputs.tsv": sha256_file(root / "raw_inputs.tsv")}:
        raise ReportValidationError(
            "CACHE_SEAL.sha256 does not bind the packet raw_inputs.tsv exactly"
        )


def read_module_metric_table(path: Path) -> dict[str, str]:
    require_plain_file(path, "module status source table")
    with path.open(encoding="utf-8", newline="") as handle:
        reader = csv.DictReader(handle, delimiter="\t")
        if tuple(reader.fieldnames or ()) != ("metric", "value"):
            raise ReportValidationError(
                f"Module status source must contain exactly metric and value columns: {path}"
            )
        rows = list(reader)
    if not rows:
        raise ReportValidationError(f"Module status source contains no metrics: {path}")
    values: dict[str, str] = {}
    for row in rows:
        metric = row.get("metric", "")
        if not metric or metric in values:
            raise ReportValidationError(
                f"Module status source contains an empty or duplicate metric: {path}"
            )
        values[metric] = row.get("value", "")
    return values


def validate_module_states(root: Path, rows: list[dict[str, str]]) -> None:
    manifest = parse_sha256_manifest(root / "artifacts.sha256")
    seen: set[tuple[str, str, str]] = set()
    for row in rows:
        key = (row["dataset"], row["case_id"], row["module"])
        if key in seen:
            raise ReportValidationError(
                "Module status matrix contains a duplicate key: " + "/".join(key)
            )
        seen.add(key)
        if row["status"] not in ALLOWED_MODULE_STATES:
            raise ReportValidationError(
                f"Invalid module state for {row['case_id']}/{row['module']}: {row['status']}"
            )
        if row["status"] == "failed":
            raise ReportValidationError(
                f"Failed module cannot appear in a passing release: "
                f"{row['case_id']}/{row['module']}"
            )
        expected_source = (
            f"observed_normalized/{row['case_id']}/{row['module']}.tsv"
        )
        if row["source_table"] != expected_source:
            raise ReportValidationError(
                f"Module status source path mismatch for {row['case_id']}/{row['module']}: "
                f"{row['source_table']!r} != {expected_source!r}"
            )
        if expected_source not in manifest:
            raise ReportValidationError(
                f"Module status source is absent from artifacts.sha256: {expected_source}"
            )
        source = packet_path(root, expected_source)
        require_plain_file(source, f"module status source {expected_source}")
        observed_hash = sha256_file(source)
        if observed_hash != manifest[expected_source]:
            raise ReportValidationError(
                f"Module status source hash mismatch for {expected_source}: "
                f"{observed_hash} != {manifest[expected_source]}"
            )

        reason = row["reason_code"].strip()
        row["reason_display"] = reason
        if row["status"] != "ok" and not reason:
            if row["status"] != "not_applicable":
                raise ReportValidationError(
                    f"Non-ok module state lacks a reason code: "
                    f"{row['case_id']}/{row['module']}"
                )
            metrics = read_module_metric_table(source)
            required = ("status", "message", "step", "read_mode", "assay_type")
            missing = [name for name in required if not metrics.get(name, "").strip()]
            if missing:
                raise ReportValidationError(
                    f"Blank-reason not_applicable source lacks required metrics for "
                    f"{row['case_id']}/{row['module']}: {', '.join(missing)}"
                )
            if metrics["status"].strip() != row["status"]:
                raise ReportValidationError(
                    f"Module status source does not match matrix status for "
                    f"{row['case_id']}/{row['module']}"
                )
            if metrics.get("reason_code", "").strip():
                raise ReportValidationError(
                    f"Blank-reason not_applicable source contains a reason_code for "
                    f"{row['case_id']}/{row['module']}"
                )
            row["reason_display"] = metrics["message"].strip()


def validate_cross_platform(
    root: Path,
    run: dict[str, object],
    rows: list[dict[str, str]],
) -> dict[str, object]:
    if not rows:
        raise ReportValidationError("Cross-platform comparison contains no evidence")
    types = Counter(row["evidence_type"] for row in rows)
    if types["normalized_scientific_table"] == 0 or types["visual_structure"] == 0:
        raise ReportValidationError(
            "Cross-platform comparison must cover scientific tables and visual structure"
        )
    for row in rows:
        if row["verdict"] != "PASS":
            raise ReportValidationError(
                f"Cross-platform comparison failed for {row['relative_path']}"
            )
        if row["evidence_type"] == "normalized_scientific_table":
            if not SHA256_RE.fullmatch(row["macos_sha256"]) or not SHA256_RE.fullmatch(
                row["ubuntu_sha256"]
            ):
                raise ReportValidationError(
                    f"Cross-platform scientific hashes are invalid: {row['relative_path']}"
                )
            if row["macos_sha256"] != row["ubuntu_sha256"]:
                raise ReportValidationError(
                    f"Cross-platform scientific hashes differ: {row['relative_path']}"
                )

    acceptance = read_json_object(
        root / "acceptance" / "cross_platform_public_reproduction.json",
        "cross-platform acceptance evidence",
    )
    expected = {
        "schema_version": EXPECTED_SCHEMA,
        "validation_profile": EXPECTED_PROFILE,
        "evidence_type": "cross_platform_public_reproduction",
        "verdict": "PASS",
        "git_commit": run["git_commit"],
        "comparison_table": "cross_platform_comparison.tsv",
    }
    for field, value in expected.items():
        if acceptance.get(field) != value:
            raise ReportValidationError(
                f"Cross-platform acceptance mismatch for {field}: "
                f"{acceptance.get(field)!r} != {value!r}"
            )
    for field in ("macos_platform", "ubuntu_platform"):
        if not isinstance(acceptance.get(field), str) or not acceptance[field]:
            raise ReportValidationError(
                f"Cross-platform acceptance lacks {field}"
            )
    return acceptance


def validate_oracles(rows: list[dict[str, str]]) -> None:
    for row in rows:
        if row["verdict"] != "PASS":
            raise ReportValidationError(
                f"Public oracle assertion is not PASS: {row['assertion_id']}"
            )


def validate_resource_usage(
    packet_root: Path,
    expected_commit: str,
    rows: list[dict[str, str]],
    raw_inputs: list[dict[str, str]],
) -> None:
    numeric_fields = (
        "wall_seconds",
        "user_cpu_seconds",
        "system_cpu_seconds",
        "max_rss_kb",
        "broad_declared_input_inventory_bytes",
        "changed_or_new_output_inventory_bytes",
    )
    measurement_ids: set[str] = set()
    case_ids: set[str] = set()
    for row in rows:
        measurement_id = row["measurement_id"].lower()
        if (
            re.fullmatch(
                r"[0-9a-f]{8}-[0-9a-f]{4}-4[0-9a-f]{3}-[89ab][0-9a-f]{3}-[0-9a-f]{12}",
                measurement_id,
                flags=re.IGNORECASE,
            )
            is None
            or measurement_id in measurement_ids
        ):
            raise ReportValidationError(
                f"Invalid or duplicate resource measurement ID: {measurement_id!r}"
            )
        measurement_ids.add(measurement_id)
        case_id = row["case_id"]
        if case_id in case_ids:
            raise ReportValidationError(f"Duplicate resource case ID: {case_id}")
        case_ids.add(case_id)
        if row["candidate_commit"] != expected_commit:
            raise ReportValidationError(
                f"Resource candidate commit mismatch for {case_id}"
            )
        expected_paths = {
            "command_path": f"commands/{case_id}.sh",
            "log_path": f"logs/{case_id}.log",
        }
        for path_field, expected_path in expected_paths.items():
            if row[path_field] != expected_path:
                raise ReportValidationError(
                    f"Resource {path_field} mismatch for {case_id}"
                )
        for digest_field in ("command_sha256", "log_sha256"):
            if SHA256_RE.fullmatch(row[digest_field]) is None:
                raise ReportValidationError(
                    f"Invalid original execution digest for {case_id}: {digest_field}"
                )
        for path_field, digest_field in (
            ("command_path", "packaged_command_sha256"),
            ("log_path", "packaged_log_sha256"),
        ):
            relative = row[path_field]
            evidence_path = packet_path(packet_root, relative)
            require_plain_file(evidence_path, relative)
            if (
                SHA256_RE.fullmatch(row[digest_field]) is None
                or sha256_file(evidence_path) != row[digest_field]
            ):
                raise ReportValidationError(
                    f"Resource {digest_field} does not bind {path_field} for {case_id}"
                )
        status = row["measurement_status"]
        if status != "measured":
            raise ReportValidationError(
                f"Required resource measurement is not measured for {case_id}: {status}"
            )
        expected_threads = RESOURCE_CASE_THREAD_SETTINGS.get(case_id)
        if row["threads"] != expected_threads:
            raise ReportValidationError(
                f"Resource thread setting mismatch for {case_id}: "
                f"{row['threads']!r} != {expected_threads!r}"
            )
        try:
            values = [float(row[field]) for field in numeric_fields]
        except ValueError as error:
            raise ReportValidationError(
                f"Invalid resource measurement for {row['case_id']}"
            ) from error
        if any(not math.isfinite(value) for value in values):
            raise ReportValidationError(
                f"Non-finite resource measurement for {row['case_id']}"
            )
        try:
            input_count = int(row["broad_declared_input_inventory_file_count"])
            output_count = int(row["changed_or_new_output_inventory_file_count"])
        except ValueError as error:
            raise ReportValidationError(
                f"Invalid resource inventory count for {row['case_id']}"
            ) from error
        if (
            str(input_count) != row["broad_declared_input_inventory_file_count"]
            or str(output_count) != row["changed_or_new_output_inventory_file_count"]
            or input_count <= 0
            or output_count < 0
            or values[0] <= 0
            or values[3] <= 0
            or values[1] < 0
            or values[2] < 0
            or values[4] <= 0
            or values[5] < 0
        ):
            raise ReportValidationError(
                f"Out-of-range resource measurement for {row['case_id']}"
            )
        if (
            row["io_measurement_method"]
            != "broad_declared_inputs_and_changed_or_new_outputs_v3"
        ):
            raise ReportValidationError(
                f"Invalid resource I/O measurement method for {row['case_id']}"
            )
        if row["broad_declared_input_inventory_scope"] != (
            "repository_root;cache_root;validation_root"
        ):
            raise ReportValidationError(
                f"Invalid broad input inventory scope for {row['case_id']}"
            )
        if row["changed_or_new_output_inventory_scope"] != (
            "cache_root;validation_root"
        ):
            raise ReportValidationError(
                f"Invalid changed/new output inventory scope for {row['case_id']}"
            )
    if case_ids != REQUIRED_RESOURCE_CASE_IDS:
        raise ReportValidationError(
            "Resource case inventory mismatch: "
            f"missing={sorted(REQUIRED_RESOURCE_CASE_IDS - case_ids)}, "
            f"unexpected={sorted(case_ids - REQUIRED_RESOURCE_CASE_IDS)}"
        )
    cache_rows = [row for row in rows if row["case_id"] == "public_cache_prepare"]
    if len(cache_rows) != 1 or cache_rows[0]["measurement_status"] != "measured":
        raise ReportValidationError(
            "Resource evidence requires one measured public_cache_prepare row"
        )
    raw_fastq_bytes = sum(int(row["bytes"]) for row in raw_inputs)
    observed = int(cache_rows[0]["changed_or_new_output_inventory_bytes"])
    if observed < raw_fastq_bytes:
        raise ReportValidationError(
            "public_cache_prepare changed/new output inventory excludes raw downloads"
        )


def validate_gm11906_source_metadata(
    packet_root: Path,
    run: dict[str, object],
    release: dict[str, object],
) -> dict[str, object]:
    path = packet_root / GM11906_SOURCE_METADATA_PACKET_PATH
    if sha256_file(path) != GM11906_SOURCE_METADATA_SHA256:
        raise ReportValidationError(
            "GM11906 official NCBI metadata snapshot SHA-256 mismatch"
        )
    payload = read_json_object(path, "GM11906 official NCBI metadata snapshot")
    records = payload.get("records")
    if (
        payload.get("schema_version") != "1.0"
        or payload.get("resource_id")
        != "gm11906_ncbi_public_source_metadata_v1"
        or payload.get("authority") != "NCBI GEO and NCBI SRA"
        or not isinstance(records, list)
    ):
        raise ReportValidationError(
            "GM11906 official NCBI metadata snapshot identity mismatch"
        )
    canonical = json.dumps(
        records, sort_keys=True, separators=(",", ":"), ensure_ascii=True
    ).encode("ascii")
    records_sha256 = hashlib.sha256(canonical).hexdigest()
    if records_sha256 != payload.get("records_sha256"):
        raise ReportValidationError(
            "GM11906 official NCBI metadata records SHA-256 mismatch"
        )
    try:
        retrieved = datetime.fromisoformat(
            str(payload["retrieval_completed_utc"]).replace("Z", "+00:00")
        )
    except (KeyError, ValueError) as error:
        raise ReportValidationError(
            "GM11906 official NCBI metadata retrieval timestamp is invalid"
        ) from error
    if retrieved.tzinfo is None or retrieved.utcoffset() is None:
        raise ReportValidationError(
            "GM11906 official NCBI metadata retrieval timestamp lacks a timezone"
        )
    by_run = {
        record.get("run_accession"): record
        for record in records
        if isinstance(record, dict)
    }
    expected = {
        "SRR10804585": ("SAMN13699362", "GSM4238454"),
        "SRR10804590": ("SAMN13699398", "GSM4238459"),
        "SRR10804657": ("SAMN13699338", "GSM4238526"),
    }
    if len(by_run) != 3 or set(by_run) != set(expected):
        raise ReportValidationError(
            "GM11906 official NCBI metadata run inventory mismatch"
        )
    for run_accession, identifiers in expected.items():
        record = by_run[run_accession]
        if (
            (record.get("biosample_accession"), record.get("geo_accession"))
            != identifiers
            or record.get("cell_line") != "GM11906"
            or record.get("library_strategy") != "ATAC-seq"
        ):
            raise ReportValidationError(
                f"GM11906 official NCBI metadata linkage mismatch for {run_accession}"
            )
    identity = {
        "path": GM11906_SOURCE_METADATA_PACKET_PATH,
        "sha256": GM11906_SOURCE_METADATA_SHA256,
        "records_sha256": records_sha256,
        "retrieval_completed_utc": payload["retrieval_completed_utc"],
        "authority": payload["authority"],
    }
    if (
        run.get("public_source_metadata") != identity
        or release.get("public_source_metadata") != identity
    ):
        raise ReportValidationError(
            "Release identity is not bound to the official NCBI metadata snapshot"
        )
    return {**identity, "records": records}


def validate_provenance_tables(root: Path, tables: dict[str, list[dict[str, str]]]) -> None:
    for row in tables["table_provenance.tsv"]:
        relative = row["packet_path"]
        artifact = packet_path(root, relative)
        require_plain_file(artifact, f"provenance table {relative}")
        if not SHA256_RE.fullmatch(row["sha256"]):
            raise ReportValidationError(f"Invalid provenance hash for {relative}")
        if sha256_file(artifact) != row["sha256"]:
            raise ReportValidationError(f"Provenance table hash mismatch: {relative}")


def validate_figures(root: Path, rows: list[dict[str, str]]) -> list[dict[str, object]]:
    validated: list[dict[str, object]] = []
    for row in rows:
        relative = row["packet_path"]
        artifact = packet_path(root, relative)
        require_plain_file(artifact, f"report-native figure {relative}")
        if row["visual_status"] != "ok":
            raise ReportValidationError(f"Figure visual status is not ok: {relative}")
        observed_hash = sha256_file(artifact)
        if observed_hash != row["sha256"]:
            raise ReportValidationError(f"Figure hash mismatch: {relative}")
        if str(artifact.stat().st_size) != row["bytes"]:
            raise ReportValidationError(f"Figure byte count mismatch: {relative}")
        try:
            with Image.open(artifact) as image:
                image.verify()
            with Image.open(artifact) as image:
                width, height = image.size
        except Exception as error:
            raise ReportValidationError(f"Unreadable packet figure: {relative}") from error
        if str(width) != row["width"] or str(height) != row["height"]:
            raise ReportValidationError(f"Figure dimensions mismatch: {relative}")
        record: dict[str, object] = dict(row)
        record.update(path=artifact, width_px=width, height_px=height)
        validated.append(record)
    if not validated:
        raise ReportValidationError("No report-native packet figures are available")
    return validated


def load_and_validate_packet(
    packet_root: Path,
    publication_path: Path | None,
    *,
    expected_release_version: str = EXPECTED_VERSION,
    expected_package_version: str = EXPECTED_PACKAGE_VERSION,
) -> ReportEvidence:
    if packet_root.is_symlink() or not packet_root.is_dir():
        raise ReportValidationError(f"Packet root is missing or symlinked: {packet_root}")
    for relative in REQUIRED_PACKET_FILES:
        if relative == "artifacts.sha256":
            require_plain_file(packet_path(packet_root, relative), relative)
        elif relative not in ALLOWED_EMPTY_COMPARISONS:
            require_plain_file(packet_path(packet_root, relative), relative)
    validate_artifact_manifest(packet_root)

    run = read_json_object(packet_root / "run.json", "run identity")
    release = read_json_object(packet_root / "release_identity.json", "release identity")
    publication = (
        read_json_object(publication_path, "GitHub publication metadata")
        if publication_path is not None
        else None
    )
    validate_identity(
        run,
        release,
        publication,
        expected_release_version=expected_release_version,
        expected_package_version=expected_package_version,
    )
    public_source_metadata = validate_gm11906_source_metadata(
        packet_root,
        run,
        release,
    )

    tables: dict[str, list[dict[str, str]]] = {}
    for name, columns in EVIDENCE_COLUMNS.items():
        tables[name] = read_tsv(packet_root / name, columns)
    for row in tables["public_data_sources.tsv"]:
        recorded = row["metadata_recorded_utc"]
        try:
            datetime.fromisoformat(recorded.replace("Z", "+00:00"))
        except ValueError as error:
            raise ReportValidationError(
                f"Public source metadata timestamp is not ISO-8601: {recorded!r}"
            ) from error
    raw_inputs = read_tsv(packet_root / "raw_inputs.tsv", RAW_INPUT_COLUMNS)

    validate_cases(run, tables["cases.tsv"])
    validate_inputs(packet_root, raw_inputs)
    validate_module_states(packet_root, tables["module_status_matrix.tsv"])
    validate_oracles(tables["oracle_assertions.tsv"])
    validate_resource_usage(
        packet_root,
        str(run["git_commit"]),
        tables["resource_usage.tsv"],
        raw_inputs,
    )
    runtime_versions = validate_runtime_versions(
        packet_root / RUNTIME_VERSIONS_PACKET_PATH,
        expected_package_version,
    )
    network_isolation = validate_network_isolation(
        packet_root / NETWORK_ISOLATION_PACKET_PATH, runtime_versions
    )
    cross = validate_cross_platform(
        packet_root, run, tables["cross_platform_comparison.tsv"]
    )
    comparison_platform = (
        cross["macos_platform"]
        if str(runtime_versions["platform_id"]).startswith("osx-")
        else cross["ubuntu_platform"]
    )
    if comparison_platform != runtime_versions["platform_id"]:
        raise ReportValidationError(
            "Cross-platform acceptance platform does not match runtime evidence"
        )
    validate_provenance_tables(packet_root, tables)
    figures = validate_figures(packet_root, tables["figure_provenance.tsv"])

    environment_text = (packet_root / "environment.txt").read_text(encoding="utf-8")
    if not environment_text.strip():
        raise ReportValidationError("environment.txt is empty")
    return ReportEvidence(
        packet_root=packet_root,
        run=run,
        release=release,
        publication=publication or {},
        tables=tables,
        raw_inputs=raw_inputs,
        environment_text=environment_text,
        runtime_versions=runtime_versions,
        network_isolation=network_isolation,
        cross_platform=cross,
        public_source_metadata=public_source_metadata,
        figures=figures,
    )


def preflight_packet(packet_root: Path) -> None:
    """Validate complete packet evidence without GitHub publication metadata."""

    run = read_json_object(packet_root / "run.json", "run identity")
    release = read_json_object(
        packet_root / "release_identity.json", "release identity"
    )
    identity = (run.get("release_version"), release.get("package_version"))
    supported = {
        ("v0.3.0", "0.3.0"),
        (EXPECTED_VERSION, EXPECTED_PACKAGE_VERSION),
    }
    if identity not in supported:
        raise ReportValidationError(
            "Packet preflight supports only internally consistent v0.3.0 or "
            f"{EXPECTED_VERSION} release identities; observed {identity!r}"
        )
    load_and_validate_packet(
        packet_root,
        None,
        expected_release_version=str(identity[0]),
        expected_package_version=str(identity[1]),
    )


def figure_priority(record: dict[str, object]) -> tuple[int, str, str]:
    name = Path(str(record["packet_path"])).name.lower()
    priorities = (
        "heteroplasmy_landscape",
        "deletion_clusters",
        "cosegregation_heatmap",
        "copy_number_proxy",
        "numt_qc",
        "variant_consequence",
        "gene_summary",
        "depth_profile",
    )
    rank = next((index for index, value in enumerate(priorities) if value in name), 99)
    return rank, str(record["case_id"]), name


def select_report_figures(
    evidence: ReportEvidence, asset_root: Path, maximum: int = 8
) -> list[FigureBlock]:
    eligible = [
        record
        for record in evidence.figures
        if Path(str(record["packet_path"])).suffix.lower() in {".png", ".jpg", ".jpeg"}
        and "montage" not in Path(str(record["packet_path"])).name.lower()
    ]
    if not eligible:
        raise ReportValidationError(
            "No raster report-native figures are eligible for DOCX embedding"
        )

    grouped: dict[str, list[dict[str, object]]] = defaultdict(list)
    for record in eligible:
        grouped[str(record["dataset"])].append(record)
    selected: list[dict[str, object]] = []
    per_dataset = max(1, maximum // max(1, len(grouped)))
    for dataset in sorted(grouped):
        selected.extend(sorted(grouped[dataset], key=figure_priority)[:per_dataset])
    if len(selected) < maximum:
        used = {str(record["packet_path"]) for record in selected}
        remainder = sorted(
            (record for record in eligible if str(record["packet_path"]) not in used),
            key=lambda record: (str(record["dataset"]), *figure_priority(record)),
        )
        selected.extend(remainder[: maximum - len(selected)])
    selected = selected[:maximum]

    asset_root.mkdir(parents=True, exist_ok=True)
    blocks: list[FigureBlock] = []
    for number, record in enumerate(selected, 1):
        source = Path(str(record["path"]))
        safe_id = SAFE_ID_RE.sub("_", str(record["figure_id"])).strip("_") or f"F{number}"
        destination = asset_root / f"{number:02d}_{safe_id}_{source.name}"
        shutil.copy2(source, destination)
        if sha256_file(destination) != record["sha256"]:
            raise ReportValidationError(
                f"Copied report figure hash mismatch: {record['packet_path']}"
            )
        title = source.stem.replace("_", " ").strip().capitalize()
        blocks.append(
            FigureBlock(
                number=number,
                dataset=str(record["dataset"]),
                case_id=str(record["case_id"]),
                title=title,
                asset_path=destination,
                asset_relative=f"{asset_root.name}/{destination.name}",
                packet_path=str(record["packet_path"]),
                sha256=str(record["sha256"]),
                width=int(record["width_px"]),
                height=int(record["height_px"]),
                source_inventory=str(record["source_inventory"]),
            )
        )
    return blocks


def compact_rows(
    rows: Iterable[dict[str, str]], columns: Sequence[str]
) -> tuple[dict[str, str], ...]:
    return tuple({column: row.get(column, "") for column in columns} for row in rows)


def inventory_rows(evidence: ReportEvidence) -> tuple[dict[str, str], ...]:
    figures: Counter[tuple[str, str]] = Counter()
    tables: Counter[tuple[str, str]] = Counter()
    for row in evidence.tables["figure_provenance.tsv"]:
        figures[(row["dataset"], row["case_id"])] += 1
    for row in evidence.tables["table_provenance.tsv"]:
        tables[(row["dataset"], row["case_id"])] += 1
    keys = sorted(set(figures) | set(tables))
    return tuple(
        {
            "dataset": dataset,
            "case_id": case_id,
            "figures": str(figures[(dataset, case_id)]),
            "tables": str(tables[(dataset, case_id)]),
        }
        for dataset, case_id in keys
    )


def default_filter_text(rows: list[dict[str, str]]) -> str:
    defaults = [row for row in rows if row.get("profile") == "default"]
    combinations = sorted(
        {
            (
                row.get("min_base_quality", ""),
                row.get("min_mapping_quality", ""),
                row.get("min_read_mean_quality", ""),
            )
            for row in defaults
        }
    )
    if len(combinations) != 1:
        raise ReportValidationError(
            "Public default filter thresholds are missing or inconsistent"
        )
    baseq, mapq, readq = combinations[0]
    return f"baseQ >= {baseq}, MAPQ >= {mapq}, read mean Q >= {readq}"


def build_report_blocks(
    evidence: ReportEvidence, figures: list[FigureBlock]
) -> list[Block]:
    tables = evidence.tables
    default_filters = default_filter_text(tables["filter_profile_results.tsv"])
    verdict_counts = evidence.run["verdict_counts"]
    platform = evidence.cross_platform
    module_counts = Counter(row["status"] for row in tables["module_status_matrix.tsv"])
    assertion_counts = Counter(row["verdict"] for row in tables["oracle_assertions.tsv"])
    runtime = evidence.runtime_versions
    isolation = evidence.network_isolation
    runtime_rows = (
        {"field": "Platform", "value": str(runtime["platform_id"])},
        {
            "field": "Operating system / architecture",
            "value": f"{runtime['system']} / {runtime['machine']}",
        },
        {"field": "Python", "value": str(runtime["python"])},
        {"field": "samtools", "value": str(runtime["samtools"])},
        {"field": "htslib", "value": str(runtime["htslib"])},
        {"field": "minimap2", "value": str(runtime["minimap2"])},
        {"field": "BWA", "value": str(runtime["bwa"])},
        {"field": "Threads", "value": str(runtime["threads"])},
        {
            "field": "Installed distribution required",
            "value": str(runtime["installed_distribution_required"]).lower(),
        },
        {
            "field": "Pinned Python packages",
            "value": "; ".join(
                f"{name}={version}"
                for name, version in sorted(runtime["packages"].items())
            ),
        },
    )
    isolation_rows = tuple(
        {"field": field, "value": isolation[field]}
        for field in (
            "platform",
            "isolation_method",
            "isolation_scope",
            "parent_loopback_control",
            "isolated_loopback_probe",
            "probe_target",
            "probe_error",
            "network_isolation_verdict",
        )
    )
    gm11906_source_rows = tuple(
        {
            "run_accession": str(record["run_accession"]),
            "experiment_accession": str(record["experiment_accession"]),
            "geo_accession": str(record["geo_accession"]),
            "biosample_accession": str(record["biosample_accession"]),
            "cell_line": str(record["cell_line"]),
            "library_strategy": str(record["library_strategy"]),
        }
        for record in evidence.public_source_metadata["records"]
    )

    blocks: list[Block] = [
        Heading(1, "Release decision"),
        Paragraph(
            f"PASS. All {evidence.run['case_count']} recorded validation cases passed, "
            f"all {sum(assertion_counts.values())} public oracle assertions passed, and "
            "the packet contains no SKIP, XFAIL, BLOCKED, or FAIL release case."
        ),
        TableBlock(
            "Release identity",
            (
                ("field", "Field"),
                ("value", "Exact value"),
            ),
            (
                {"field": "Version/tag", "value": str(evidence.run["release_version"])},
                {"field": "Commit", "value": str(evidence.run["git_commit"])},
                {"field": "Repository", "value": str(evidence.run["repository"])},
                {
                    "field": "GitHub Actions run",
                    "value": str(evidence.run["github_actions_run_id"]),
                },
                {
                    "field": "GitHub release",
                    "value": str(evidence.publication["github_release_url"]),
                },
                {
                    "field": "Publication evidence scope",
                    "value": (
                        "read-only prepublication tag/repository identity; final asset "
                        "publication is verified separately in github_publication.json"
                    ),
                },
                {"field": "Packet schema", "value": EXPECTED_SCHEMA},
                {"field": "Validation profile", "value": EXPECTED_PROFILE},
            ),
        ),
        Heading(1, "Scope and claim boundary"),
        Paragraph(
            "This is a release-validation report for a reproducible, mode-gated mtDNA "
            "reporting workflow/resource. It documents execution contracts, fixed-input "
            "repeatability, public marker representation, descriptive filter dependence, "
            "and release provenance. It does not establish diagnostic performance, "
            "sensitivity or specificity, a limit of detection, deletion-calling accuracy, "
            "absolute mtDNA copy number, formal NUMT classification, modality equivalence, "
            "or population generalizability."
        ),
        Paragraph(
            "This report is generated before GitHub release creation because the report "
            "is itself a hashed release asset. The separate post-publication receipt "
            "github_publication.json verifies the uploaded asset inventory, authenticated "
            "redownload hashes, annotated tag, and final hosting-protection state without "
            "making this report self-referential. Native GitHub release immutability must "
            "be enabled and confirmed before draft creation; publication fails closed if "
            "enablement or the post-publication immutability query does not succeed."
        ),
        TableBlock(
            "Bounded claim-to-evidence mapping",
            (
                ("claim_id", "ID"),
                ("bounded_claim", "Bounded claim"),
                ("evidence", "Evidence"),
                ("limitation", "Limitation"),
            ),
            tuple(tables["claim_evidence_matrix.tsv"]),
        ),
        Heading(1, "Five v0.3.0 scientific corrections"),
        Heading(2, "1. Filtered alternate-allele observation counting"),
        Paragraph(
            "The release uses one observation engine for candidate selection and "
            "co-occurrence calculations. Callable depth is the sum of passing A, C, G, "
            f"and T observations. The public default profile is {default_filters}. "
            "These are reporting defaults rather than clinically calibrated thresholds."
        ),
        Formula("AF_alt = N_alt / (N_A + N_C + N_G + N_T)"),
        Paragraph(
            "Required identities are N_alt = N_alt,forward + N_alt,reverse and callable "
            "depth = N_A + N_C + N_G + N_T. The no-cap setting prevents a hidden 8,000-read "
            "pileup truncation."
        ),
        Heading(2, "2. Explicit mvTool network control"),
        Paragraph(
            "mvTool operation is mode-gated as disabled, fixture, or network. Disabled is "
            "the default and produces a deterministic not_configured result without an "
            "HTTP request. Fixture mode supports deterministic testing. A requested network "
            "failure is reported as unavailable and does not fabricate an annotation."
        ),
        Heading(2, "3. Standalone input contract"),
        Paragraph(
            "The core contract requires WORK_ROOT, RUN_NAME, SAMPLE_ID, REF_FASTA, "
            "SOURCE_ALIGN_FILE, and MT_CONTIG. Alignment type is inferred from BAM or CRAM, "
            "MT length is read from the FASTA index, and optional variant or methylation "
            "sidecars resolve by explicit path, legacy discovery, then absence. Missing "
            "optional sidecars do not fail core reporting."
        ),
        Heading(2, "4. Experimental mt:nuclear depth proxy"),
        Paragraph(
            "The copy-number layer remains a within-sample depth proxy. Successfully "
            "measured zero-depth nuclear windows contribute to the nuclear mean; an all-zero "
            "or otherwise invalid denominator yields not_evaluable rather than zero. No "
            "diploid-cell copy-number claim is made."
        ),
        Formula("R_mt:nuclear = mean(D_mt) / mean(D_nuclear windows)"),
        Heading(2, "5. Reference scope, alignment ambiguity, and BED convention"),
        Paragraph(
            "REFERENCE_SCOPE distinguishes mt-only, whole-genome, and custom references. "
            "Under mt-only or custom scope, raw alignment-span and MAPQ metrics remain "
            "available but NUMT interpretation is not_evaluable. The mitochondrial interval "
            "uses a zero-based, half-open BED record."
        ),
        Formula("BED interval = MT_CONTIG <TAB> 0 <TAB> MT_LENGTH"),
        Heading(1, "Additional bounded algorithms"),
        Paragraph(
            "Deletion support is a read-alignment screening metric, not a validated deletion "
            "caller. Its primary-read support fraction uses supporting query names that have "
            "a retained primary mtDNA alignment over unique retained primary mtDNA query "
            "names, preserving a value in [0,1]. Co-occurrence uses the identical passing "
            "read-observation set used by candidate selection."
        ),
        Formula(
            "f_deletion,primary = N_supporting names with primary mt alignment / "
            "N_unique primary mt query names"
        ),
        Formula("J(A,B) = |R_A intersection R_B| / |R_A union R_B|"),
        Heading(1, "Frozen public inputs and provenance"),
        Paragraph(
            f"The sealed input manifest contains {len(evidence.raw_inputs)} public FASTQ "
            "objects. Raw reads are not redistributed in the release packet. Every filename "
            "and SHA-256 below is cross-checked against inputs.sha256 before this report is "
            "created."
        ),
        TableBlock(
            "Exact public FASTQ inventory",
            (
                ("dataset_id", "Dataset"),
                ("run_accession", "Run"),
                ("filename", "Filename"),
                ("bytes", "Bytes"),
                ("fastq_records", "Records"),
                ("sha256", "SHA-256"),
            ),
            tuple(evidence.raw_inputs),
        ),
        TableBlock(
            "Exact public FASTQ source locations",
            (
                ("filename", "Filename"),
                ("source_record_url", "Metadata record"),
                ("url", "FASTQ URL"),
            ),
            tuple(evidence.raw_inputs),
        ),
        TableBlock(
            "Public source metadata",
            (
                ("dataset", "Dataset"),
                ("run_accession", "Run"),
                ("study_accession", "Study"),
                ("sample_accession", "Sample"),
                ("platform", "Platform"),
                ("library_strategy", "Library"),
                ("metadata_recorded_utc", "metadata_recorded_utc"),
                ("role", "Validation role"),
            ),
            tuple(tables["public_data_sources.tsv"]),
        ),
        Paragraph(
            "The three GM11906 run-to-sample relationships are bound to a tracked "
            "official NCBI GEO/SRA snapshot captured at "
            f"{evidence.public_source_metadata['retrieval_completed_utc']}. The packet "
            "verifies the snapshot at SHA-256 "
            f"{evidence.public_source_metadata['sha256']}; no live metadata lookup is "
            "used during offline validation."
        ),
        TableBlock(
            "Official NCBI GM11906 accession linkage",
            (
                ("run_accession", "Run"),
                ("experiment_accession", "Experiment"),
                ("geo_accession", "GEO sample"),
                ("biosample_accession", "BioSample"),
                ("cell_line", "Cell line"),
                ("library_strategy", "Library"),
            ),
            gm11906_source_rows,
        ),
        Heading(1, "Execution environment"),
        Paragraph(
            "The packet records the release identity, platform, software versions, and "
            "execution controls. The verbatim environment record follows so that values are "
            "not paraphrased or reconstructed."
        ),
        TableBlock(
            "Validated installed runtime",
            (("field", "Field"), ("value", "Recorded value")),
            runtime_rows,
        ),
        Paragraph(
            "The public-data matrix ran in one operating-system-isolated process tree. "
            "A parent loopback listener was reachable before isolation, while the probe "
            "from the isolated child was blocked before the matrix began."
        ),
        TableBlock(
            "OS-level network-isolation evidence",
            (("field", "Field"), ("value", "Recorded value")),
            isolation_rows,
        ),
        CodeBlock(evidence.environment_text.rstrip()),
        Heading(1, "Validation cases"),
        Paragraph(
            f"Recorded verdicts: {', '.join(f'{key}={value}' for key, value in sorted(verdict_counts.items()))}. "
            "Every listed case had both input and expected evidence available."
        ),
        TableBlock(
            "Complete validation case ledger",
            (
                ("case_id", "Case"),
                ("category", "Category"),
                ("verdict", "Verdict"),
                ("detail", "Evidence detail"),
            ),
            tuple(tables["cases.tsv"]),
        ),
        Heading(1, "Public-data oracle results"),
        Paragraph(
            "The following descriptive results are read directly from the corrected release "
            "packet. A PASS means equality with the reviewed fixed-input oracle; it is not a "
            "clinical sensitivity or accuracy estimate."
        ),
        TableBlock(
            "Filter-profile results",
            (
                ("dataset", "Dataset"),
                ("profile", "Profile"),
                ("candidate_sites", "Candidates"),
                ("accepted_observations", "Accepted"),
                ("excluded_observations", "Excluded"),
                ("m8344_A_G_alt_allele_fraction", "m.8344 AF"),
            ),
            compact_rows(
                tables["filter_profile_results.tsv"],
                (
                    "dataset",
                    "profile",
                    "candidate_sites",
                    "accepted_observations",
                    "excluded_observations",
                    "m8344_A_G_alt_allele_fraction",
                ),
            ),
        ),
        TableBlock(
            "Public report output inventory",
            (
                ("dataset", "Dataset"),
                ("profile", "Profile"),
                ("summary_tsv_count", "Summary TSV"),
                ("html_count", "HTML"),
                ("png_count", "PNG"),
            ),
            compact_rows(
                tables["public_validation_oracle_v0.3.0.tsv"],
                (
                    "dataset",
                    "profile",
                    "summary_tsv_count",
                    "html_count",
                    "png_count",
                ),
            ),
        ),
        TableBlock(
            "Exact oracle assertions",
            (
                ("assertion_id", "Assertion"),
                ("verdict", "Verdict"),
                ("expected", "Expected"),
                ("observed", "Observed"),
                ("detail", "Detail"),
            ),
            tuple(tables["oracle_assertions.tsv"]),
        ),
        TableBlock(
            "Release-result handoff values",
            (
                ("result_id", "ID"),
                ("dataset", "Dataset"),
                ("metric", "Metric"),
                ("value", "Value"),
                ("unit", "Unit"),
                ("claim_boundary", "Boundary"),
            ),
            tuple(tables["manuscript_handoff.tsv"]),
        ),
        Heading(1, "Module-state interpretation"),
        Paragraph(
            "Module states describe workflow applicability and availability, not biological "
            f"normality. State counts are {', '.join(f'{key}={value}' for key, value in sorted(module_counts.items()))}. "
            "Expected not_configured, not_applicable, or not_evaluable states can coexist "
            "with a passing workflow only when prescribed by the mode and reference scope."
        ),
        TableBlock(
            "Module status matrix",
            (
                ("dataset", "Dataset"),
                ("case_id", "Case"),
                ("module", "Module"),
                ("status", "State"),
                ("reason_display", "Reason"),
                ("source_table", "Evidence table"),
            ),
            tuple(tables["module_status_matrix.tsv"]),
        ),
        Heading(1, "Repeatability and cross-platform reproduction"),
        Paragraph(
            f"The packet records independent macOS ({platform['macos_platform']}) and "
            f"Ubuntu ({platform['ubuntu_platform']}) public-data executions at the exact "
            "release commit. Each visual inventory is first bound to its platform's actual "
            "HTML/PNG artifacts, including byte count, SHA-256, and decoded PNG dimensions. "
            "Normalized scientific TSVs are required to match exactly. PNG byte hashes are "
            "not cross-platform gates; visual path, type, dimensions, and integrity are "
            "compared instead."
        ),
        TableBlock(
            "Cross-platform comparison",
            (
                ("evidence_type", "Evidence type"),
                ("relative_path", "Relative path"),
                ("verdict", "Verdict"),
                ("comparison", "Comparison rule"),
            ),
            tuple(tables["cross_platform_comparison.tsv"]),
        ),
        Heading(1, "Evidence verification trust boundary"),
        Paragraph(
            "Release verification is ordered deliberately. First, the validation ZIP is "
            "compared with an expected SHA-256 supplied outside that ZIP, using the "
            "release identity during assembly and the published SHA256SUMS manifest after "
            "publication. Only after that comparison passes is the archive safely "
            "extracted and its verify_bundle.sh executed. The packet verifier establishes "
            "internal inventory and semantic consistency; it does not authenticate a ZIP "
            "against coordinated replacement of the archive and all internal hashes. The "
            "external digest source therefore remains the trust anchor and must itself be "
            "obtained through the exact tag or immutable GitHub release record."
        ),
        Heading(1, "Resource measurements"),
        Paragraph(
            "Resource values are observational measurements for the recorded hardware and "
            "inputs. Byte values are inventories, not operating-system I/O counters: the "
            "input count and byte values are the broad pre-command inventory of the "
            "repository, cache, and validation roots, while the output count and byte "
            "values describe files created or changed under the cache and validation "
            "roots. Every required row is measured, finite, and has positive wall time, "
            "peak RSS, input count, and input bytes. They are not generalized "
            "performance benchmarks. The thread column records the configured workflow "
            "setting for each case; orchestration or mixed test suites are labeled "
            "not_applicable or mixed rather than assigned an artificial thread count. "
            "Original execution hashes are retained separately from hashes of portable, "
            "path-sanitized command and log copies."
        ),
        TableBlock(
            "Recorded resource usage",
            (
                ("case_id", "Case"),
                ("wall_seconds", "Wall s"),
                ("user_cpu_seconds", "User CPU s"),
                ("system_cpu_seconds", "System CPU s"),
                ("max_rss_kb", "Peak RSS KB"),
                (
                    "broad_declared_input_inventory_file_count",
                    "Broad declared input file count",
                ),
                (
                    "broad_declared_input_inventory_bytes",
                    "Broad declared input inventory bytes",
                ),
                (
                    "changed_or_new_output_inventory_file_count",
                    "Changed/new output file count",
                ),
                (
                    "changed_or_new_output_inventory_bytes",
                    "Changed/new output inventory bytes",
                ),
                ("threads", "Configured workflow threads"),
                ("platform", "Platform"),
                ("measurement_status", "Status"),
            ),
            tuple(tables["resource_usage.tsv"]),
        ),
        Heading(1, "Output inventory and evidence provenance"),
        TableBlock(
            "Report-native artifact inventory",
            (
                ("dataset", "Dataset"),
                ("case_id", "Case"),
                ("figures", "Figures"),
                ("tables", "Tables"),
            ),
            inventory_rows(evidence),
        ),
        Paragraph(
            "All embedded figures below are copied from the exact-commit packet after "
            "SHA-256, byte-count, image-integrity, and dimension checks. No simplified "
            "replacement chart is generated by this report builder."
        ),
        Heading(1, "Report-native figures"),
    ]

    for index, figure in enumerate(figures):
        blocks.append(figure)
        if index != len(figures) - 1:
            blocks.append(PageBreak())

    figure_rows = tuple(
        {
            "figure": f"Figure {figure.number}",
            "dataset": figure.dataset,
            "case_id": figure.case_id,
            "packet_path": figure.packet_path,
            "sha256": figure.sha256,
            "dimensions": f"{figure.width} x {figure.height}",
            "source_inventory": figure.source_inventory,
        }
        for figure in figures
    )
    blocks.extend(
        [
            Heading(1, "Figure provenance ledger"),
            TableBlock(
                "Embedded figure provenance",
                (
                    ("figure", "Figure"),
                    ("dataset", "Dataset"),
                    ("case_id", "Case"),
                    ("packet_path", "Packet path"),
                    ("sha256", "SHA-256"),
                    ("dimensions", "Pixels"),
                ),
                figure_rows,
            ),
            Heading(1, "Limitations"),
            TableBlock(
                "Release limitations",
                (
                    ("limitation_id", "ID"),
                    ("scope", "Scope"),
                    ("limitation", "Limitation"),
                    ("release_effect", "Effect on claims"),
                ),
                tuple(tables["limitations.tsv"]),
            ),
            Heading(1, "Audit handoff"),
            Paragraph(
                "The machine-verifiable packet remains the authority for commands, logs, "
                "normalized outputs, source manifests, expected values, package artifacts, "
                "and hashes. This human-readable report is a deterministic view over that "
                "evidence. PDF conversion is an optional delivery handoff and does not "
                "replace rendered-page inspection of the DOCX."
            ),
        ]
    )
    return blocks


def markdown_escape(value: object) -> str:
    return str(value).replace("|", "\\|").replace("\n", "<br>")


def markdown_table(block: TableBlock) -> list[str]:
    lines = [f"**{block.title}.**", ""]
    lines.append("| " + " | ".join(label for _, label in block.columns) + " |")
    lines.append("| " + " | ".join("---" for _ in block.columns) + " |")
    for row in block.rows:
        lines.append(
            "| "
            + " | ".join(markdown_escape(row.get(key, "")) for key, _ in block.columns)
            + " |"
        )
    lines.append("")
    return lines


def render_markdown(
    evidence: ReportEvidence, blocks: list[Block], output: Path
) -> None:
    lines = [
        f"# {OUTPUT_STEM.replace('_', ' ')}",
        "",
        f"**Release:** `{evidence.run['release_version']}`  ",
        f"**Exact commit:** `{evidence.run['git_commit']}`  ",
        f"**Repository:** {evidence.run['repository']}  ",
        f"**Validation profile:** `{evidence.run['validation_profile']}`  ",
        f"**Packet generated:** `{evidence.run['generated_utc']}`",
        "",
    ]
    for block in blocks:
        if isinstance(block, Heading):
            lines.extend(["#" * (block.level + 1) + f" {block.text}", ""])
        elif isinstance(block, Paragraph):
            lines.extend([block.text, ""])
        elif isinstance(block, Formula):
            lines.extend([f"```text\n{block.text}\n```", ""])
        elif isinstance(block, CodeBlock):
            lines.extend([f"```text\n{block.text}\n```", ""])
        elif isinstance(block, TableBlock):
            lines.extend(markdown_table(block))
        elif isinstance(block, FigureBlock):
            caption = (
                f"Figure {block.number}. {block.title} for {block.dataset} "
                f"({block.case_id}). Packet source `{block.packet_path}`; SHA-256 "
                f"`{block.sha256}`; {block.width} x {block.height} pixels; release commit "
                f"`{evidence.run['git_commit']}`."
            )
            lines.extend(
                [
                    f"![{markdown_escape(caption)}]({block.asset_relative})",
                    "",
                    f"*{caption}*",
                    "",
                ]
            )
        elif isinstance(block, PageBreak):
            lines.extend(["<div style=\"page-break-after: always;\"></div>", ""])
    output.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[int]) -> None:
    if sum(widths) != 9360:
        raise ReportValidationError(f"DOCX table widths must sum to 9360 DXA: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def table_widths(columns: Sequence[tuple[str, str]]) -> tuple[int, ...]:
    weights = []
    for key, _ in columns:
        if "sha256" in key:
            weights.append(3.0)
        elif key in {"detail", "limitation", "release_effect", "bounded_claim", "evidence", "comparison"}:
            weights.append(2.6)
        elif key in {"packet_path", "relative_path", "source_table", "claim_boundary", "filename"}:
            weights.append(2.1)
        elif key in {"verdict", "status", "threads", "figures", "tables", "unit"}:
            weights.append(0.8)
        else:
            weights.append(1.25)
    total = sum(weights)
    raw = [int(9360 * weight / total) for weight in weights]
    raw[-1] += 9360 - sum(raw)
    return tuple(raw)


def docx_cell_text(value: object, key: str) -> str:
    text = str(value)
    if "sha256" in key and SHA256_RE.fullmatch(text):
        return text[:32] + "\n" + text[32:]
    if key in {"git_commit", "commit"} and COMMIT_RE.fullmatch(text):
        return text[:20] + "\n" + text[20:]
    return text


def add_docx_table(document: Document, block: TableBlock) -> None:
    caption = document.add_paragraph()
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_next = True
    set_run_font(caption.add_run(block.title + "."), size=9, bold=True, color="1F4D78")

    table = document.add_table(rows=1, cols=len(block.columns))
    table.style = "Table Grid"
    widths = table_widths(block.columns)
    header = table.rows[0]
    repeat_table_header(header)
    for cell, (_, label) in zip(header.cells, block.columns, strict=True):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "E8EEF5")
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(paragraph.add_run(label), size=8, bold=True, color="0B2545")

    for source_row in block.rows:
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        for cell, (key, _) in zip(cells, block.columns, strict=True):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(docx_cell_text(source_row.get(key, ""), key))
            set_run_font(
                run,
                name="Consolas" if "sha256" in key else "Calibri",
                size=7.2 if len(block.columns) >= 6 else 8,
            )
    set_table_geometry(table, widths)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))


def configure_document(document: Document, evidence: ReportEvidence) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False
    document.settings.odd_and_even_pages_header_footer = True

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 24, "0B2545", 0, 8),
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for header in (section.header.paragraphs[0], section.even_page_header.paragraphs[0]):
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(
            header.add_run("MitoOverview | v0.3.1 release validation"),
            size=8.5,
            color="6B7280",
        )
    for footer in (section.footer.paragraphs[0], section.even_page_footer.paragraphs[0]):
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_run_font(footer.add_run("Exact commit report | "), size=8, color="6B7280")
        add_page_field(footer)

    document.core_properties.title = OUTPUT_STEM.replace("_", " ")
    document.core_properties.subject = "GitHub release validation evidence"
    document.core_properties.author = "MitoOverview release validation"
    document.core_properties.keywords = (
        "mitochondrial DNA; bioinformatics; release validation; reproducibility"
    )
    generated = str(evidence.run.get("generated_utc", ""))
    try:
        timestamp = datetime.fromisoformat(generated.replace("Z", "+00:00"))
        document.core_properties.created = timestamp.replace(tzinfo=None)
        document.core_properties.modified = timestamp.replace(tzinfo=None)
    except ValueError:
        pass


def add_formula(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.right_indent = Inches(0.2)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F4F7")
    properties.append(shading)
    set_run_font(paragraph.add_run(text), name="Consolas", size=9.5, color="0B2545")


def add_code_block(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.0
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F7F7F7")
    properties.append(shading)
    set_run_font(paragraph.add_run(text), name="Consolas", size=7.5, color="222222")


def add_figure(document: Document, figure: FigureBlock, commit: str) -> None:
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.space_after = Pt(4)
    picture_paragraph.paragraph_format.keep_with_next = True
    width_inches = 6.15
    height_inches = width_inches * figure.height / figure.width
    if height_inches > 6.6:
        height_inches = 6.6
        width_inches = height_inches * figure.width / figure.height
    run = picture_paragraph.add_run()
    run.add_picture(
        str(figure.asset_path), width=Inches(width_inches), height=Inches(height_inches)
    )

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(6)
    caption_text = (
        f"Figure {figure.number}. {figure.title} for {figure.dataset} "
        f"({figure.case_id}). Packet source {figure.packet_path}; SHA-256 "
        f"{figure.sha256}; {figure.width} x {figure.height} pixels; release commit {commit}."
    )
    set_run_font(caption.add_run(caption_text), size=8.5, italic=True, color="4B5563")


def render_docx(evidence: ReportEvidence, blocks: list[Block], output: Path) -> None:
    document = Document()
    configure_document(document, evidence)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(
        title.add_run("MitoOverview v0.3.1"),
        size=24,
        color="0B2545",
    )
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    set_run_font(
        subtitle.add_run("GitHub release and clean-room validation report"),
        size=13,
        bold=True,
        color="2E74B5",
    )
    metadata = (
        ("Release", str(evidence.run["release_version"])),
        ("Exact commit", str(evidence.run["git_commit"])),
        ("Repository", str(evidence.run["repository"])),
        ("Packet generated", str(evidence.run["generated_utc"])),
    )
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_run_font(paragraph.add_run(label + ": "), size=9.5, bold=True)
        set_run_font(
            paragraph.add_run(value),
            name="Consolas" if label == "Exact commit" else "Calibri",
            size=9.2,
        )

    for block in blocks:
        if isinstance(block, Heading):
            level = min(max(block.level, 1), 3)
            document.add_heading(block.text, level=level)
        elif isinstance(block, Paragraph):
            document.add_paragraph(block.text)
        elif isinstance(block, Formula):
            add_formula(document, block.text)
        elif isinstance(block, CodeBlock):
            add_code_block(document, block.text)
        elif isinstance(block, TableBlock):
            add_docx_table(document, block)
        elif isinstance(block, FigureBlock):
            add_figure(document, block, str(evidence.run["git_commit"]))
        elif isinstance(block, PageBreak):
            document.add_page_break()
    document.save(output)


def write_figure_manifest(figures: list[FigureBlock], path: Path) -> None:
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle, delimiter="\t", lineterminator="\n")
        writer.writerow(
            (
                "figure_number",
                "dataset",
                "case_id",
                "report_asset",
                "packet_path",
                "sha256",
                "width_px",
                "height_px",
                "source_inventory",
            )
        )
        for figure in figures:
            writer.writerow(
                (
                    figure.number,
                    figure.dataset,
                    figure.case_id,
                    figure.asset_relative,
                    figure.packet_path,
                    figure.sha256,
                    figure.width,
                    figure.height,
                    figure.source_inventory,
                )
            )


def write_build_provenance(
    evidence: ReportEvidence,
    publication_json: Path,
    report_md: Path,
    report_docx: Path,
    report_pdf: Path | None,
    figures: list[FigureBlock],
    asset_root: Path,
) -> Path:
    """Bind generated report sources and copied figures to validated packet evidence."""

    manifest = asset_root / "figure_manifest.tsv"
    outputs = {
        "markdown": file_record(report_md),
        "docx": file_record(report_docx),
    }
    if report_pdf is not None:
        outputs["pdf"] = file_record(report_pdf)

    figure_rows: list[dict[str, object]] = []
    for figure in figures:
        asset = asset_root / Path(figure.asset_relative).name
        record = file_record(asset, name=figure.asset_relative)
        if record["sha256"] != figure.sha256:
            raise ReportValidationError(
                f"Report figure changed before provenance capture: {figure.asset_relative}"
            )
        figure_rows.append(
            {
                **record,
                "figure_number": figure.number,
                "dataset": figure.dataset,
                "case_id": figure.case_id,
                "packet_path": figure.packet_path,
                "packet_sha256": figure.sha256,
                "width_px": figure.width,
                "height_px": figure.height,
                "source_inventory": figure.source_inventory,
            }
        )

    packet_files = {
        name: file_record(evidence.packet_root / name, name=name)
        for name in ("run.json", "release_identity.json", "artifacts.sha256")
    }
    payload = {
        "schema_version": "1.0",
        "provenance_type": "mito_overview_release_report_build",
        "repository": evidence.run["repository"],
        "release_version": evidence.run["release_version"],
        "release_tag": evidence.publication["release_tag"],
        "git_commit": evidence.run["git_commit"],
        "validation_profile": evidence.run["validation_profile"],
        "packet_identity": packet_files,
        "publication_input": file_record(
            publication_json, name=publication_json.name
        ),
        "report_outputs": outputs,
        "figure_manifest": file_record(
            manifest, name=f"{asset_root.name}/figure_manifest.tsv"
        ),
        "figures": figure_rows,
        "pdf_included": report_pdf is not None,
        "rendered_page_qa_required": True,
    }
    destination = asset_root / BUILD_PROVENANCE_NAME
    destination.write_text(
        json.dumps(payload, indent=2, sort_keys=True) + "\n", encoding="utf-8"
    )
    return destination


def emit_pdf(docx_path: Path, pdf_path: Path) -> None:
    soffice = shutil.which("soffice")
    if soffice is None:
        mac_soffice = Path(
            "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        )
        if mac_soffice.is_file():
            soffice = str(mac_soffice)
    if soffice is None:
        raise ReportValidationError(
            "--emit-pdf requires LibreOffice/soffice; use the Documents render workflow "
            "for the final PDF and page-image QA"
        )
    with tempfile.TemporaryDirectory(prefix="mito-report-pdf-") as temporary:
        temporary_path = Path(temporary)
        result = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(temporary_path),
                str(docx_path),
            ],
            check=False,
            capture_output=True,
            text=True,
        )
        generated = temporary_path / (docx_path.stem + ".pdf")
        if result.returncode != 0 or not generated.is_file() or generated.stat().st_size == 0:
            raise ReportValidationError(
                "LibreOffice PDF handoff failed: "
                + (result.stderr.strip() or result.stdout.strip() or "no PDF produced")
            )
        shutil.copy2(generated, pdf_path)


def generate_report(
    packet_root: Path,
    publication_json: Path,
    output_dir: Path,
    *,
    overwrite: bool = False,
    include_pdf: bool = False,
) -> dict[str, Path]:
    evidence = load_and_validate_packet(packet_root, publication_json)
    resolved_packet = packet_root.resolve()
    resolved_output = output_dir.resolve()
    try:
        resolved_output.relative_to(resolved_packet)
    except ValueError:
        pass
    else:
        raise ReportValidationError("Report output directory must be outside the packet root")

    output_dir.mkdir(parents=True, exist_ok=True)
    md_path = output_dir / f"{OUTPUT_STEM}.md"
    docx_path = output_dir / f"{OUTPUT_STEM}.docx"
    pdf_path = output_dir / f"{OUTPUT_STEM}.pdf"
    asset_path = output_dir / f"{OUTPUT_STEM}_assets"
    targets = [md_path, docx_path, asset_path]
    if include_pdf:
        targets.append(pdf_path)
    existing = [path for path in targets if path.exists()]
    if existing and not overwrite:
        raise ReportValidationError(
            "Report output already exists; pass --overwrite to replace: "
            + ", ".join(str(path) for path in existing)
        )

    with tempfile.TemporaryDirectory(prefix=".mito-report-", dir=output_dir) as temporary:
        stage = Path(temporary)
        stage_assets = stage / asset_path.name
        figures = select_report_figures(evidence, stage_assets)
        blocks = build_report_blocks(evidence, figures)
        stage_md = stage / md_path.name
        stage_docx = stage / docx_path.name
        render_markdown(evidence, blocks, stage_md)
        render_docx(evidence, blocks, stage_docx)
        write_figure_manifest(figures, stage_assets / "figure_manifest.tsv")
        stage_pdf = stage / pdf_path.name
        if include_pdf:
            emit_pdf(stage_docx, stage_pdf)
        stage_provenance = write_build_provenance(
            evidence,
            publication_json,
            stage_md,
            stage_docx,
            stage_pdf if include_pdf else None,
            figures,
            stage_assets,
        )

        for path in existing:
            if path.is_dir():
                shutil.rmtree(path)
            else:
                path.unlink()
        os.replace(stage_assets, asset_path)
        os.replace(stage_md, md_path)
        os.replace(stage_docx, docx_path)
        if include_pdf:
            os.replace(stage_pdf, pdf_path)

    outputs = {
        "markdown": md_path,
        "docx": docx_path,
        "assets": asset_path,
        "build_provenance": asset_path / stage_provenance.name,
    }
    if include_pdf:
        outputs["pdf"] = pdf_path
    return outputs


def main() -> None:
    args = parse_args()
    try:
        if args.preflight_packet:
            preflight_packet(args.packet_root)
            print(f"[PREFLIGHT] PASS: {args.packet_root}")
            return
        outputs = generate_report(
            args.packet_root,
            args.publication_json,
            args.output_dir,
            overwrite=args.overwrite,
            include_pdf=args.emit_pdf,
        )
    except ReportValidationError as error:
        raise SystemExit(f"[ERROR] {error}") from error
    for label, path in outputs.items():
        print(f"[REPORT] {label}: {path}")


if __name__ == "__main__":
    main()
